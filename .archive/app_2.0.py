# app.py
# Final Version: Multi-Provider, Two-Model Selection, Permissive Google Safety, Default Temperature, No Overwrite
# Revision: Include custom instructions in Step 2 (Global Context) and update UI guidance for name spellings.
# Revision: Integrated Topic-Based Chunking

import streamlit as st
import os
import io  # For handling file uploads and downloads in memory
from dotenv import load_dotenv
import string # For fallback text cleaning in DOCX generation
import traceback # For detailed error logging
import requests # For OpenRouter model fetching

# --- Set page config FIRST ---
st.set_page_config(layout="wide", page_title="Deposition Summarizer")

# --- Import additional modules for enhancements ---
import logging
import platform
import time
import hashlib
from datetime import datetime
# Import critical patch functions
# Functions previously imported from critical_fixes_patch are now defined below

# --- !!! DEFINE OUTPUT DIRECTORY HERE !!! ---
# Use a raw string (r"...") or escaped backslashes for Windows paths
OUTPUT_DIRECTORY = r"D:\OneDrive\OneDrive - Barnes Maloney PLLC\Documents\Depo Summaries"
# OUTPUT_DIRECTORY = "D:\\OneDrive\\OneDrive - Barnes Maloney PLLC\\Documents\\Depo Summaries" # Alternative

# --- Application Metadata for OpenRouter ---
# Optional: Replace with your actual site URL and app name if you deploy publicly
YOUR_SITE_URL = "http://localhost:8501" # Default Streamlit URL
YOUR_APP_NAME = "Deposition Summarizer"
# ---

# --- Setup Logging ---
log_dir = os.path.join(OUTPUT_DIRECTORY, "logs")
os.makedirs(log_dir, exist_ok=True)
log_file = os.path.join(log_dir, f"depo_summarizer_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file),
        logging.StreamHandler()  # Also log to console
    ]
)
logger = logging.getLogger(__name__)
logger.info(f"Deposition Summarizer started. Output directory: {OUTPUT_DIRECTORY}")

# --- Load API Keys from .env ---
load_dotenv()
api_keys = {
    "OpenAI": os.getenv("OPENAI_API_KEY"),
    "Anthropic": os.getenv("ANTHROPIC_API_KEY"),
    "Google": os.getenv("GOOGLE_API_KEY"),
    "OpenRouter": os.getenv("OPENROUTER_API_KEY"),
}

# --- Import Libraries ---
# LLM Libraries
from openai import OpenAI, OpenAIError
from anthropic import Anthropic, APIError as AnthropicAPIError
import google.generativeai as genai
from google.api_core import exceptions as GoogleAPICoreExceptions
from google.generativeai.types import HarmCategory, HarmBlockThreshold

# File Parsing Libraries
import docx # python-docx library
import fitz # PyMuPDF library, often imported as fitz

# Text Processing (Chunking)
from langchain_text_splitters import RecursiveCharacterTextSplitter

# <<< NEW IMPORTS FOR TOPIC-BASED CHUNKING >>>
import re
from typing import List, Tuple, Dict, Callable # Added Tuple, Dict, Callable
from dataclasses import dataclass
from enhanced_topic_chunker_fixed import EnhancedTopicBasedChunker, TopicSegment as EnhancedTopicSegment
# <<< END NEW IMPORTS >>>


# --- Initialize LLM Clients (if keys exist) ---
# This dictionary will hold initialized clients for available providers
clients = {}
client_load_status = {} # To show status in sidebar

# Initialize OpenAI Client
if api_keys["OpenAI"]:
    try:
        clients["OpenAI"] = OpenAI(api_key=api_keys["OpenAI"])
        client_load_status["OpenAI"] = "✅ Initialized"
    except OpenAIError as e:
        client_load_status["OpenAI"] = f"❌ OpenAI Error: {e}"
    except Exception as e:
        client_load_status["OpenAI"] = f"❌ Unexpected Error: {e}"
else:
    client_load_status["OpenAI"] = "⚠️ Key Missing"

# Initialize Anthropic (Claude) Client
if api_keys["Anthropic"]:
    try:
        clients["Anthropic"] = Anthropic(api_key=api_keys["Anthropic"])
        client_load_status["Anthropic"] = "✅ Initialized"
    except AnthropicAPIError as e:
        client_load_status["Anthropic"] = f"❌ Anthropic Error: {e}"
    except Exception as e:
        client_load_status["Anthropic"] = f"❌ Unexpected Error: {e}"
else:
    client_load_status["Anthropic"] = "⚠️ Key Missing"

# Initialize Google (Gemini) Client
if api_keys["Google"]:
    try:
        # Configure the genai library with the API key
        genai.configure(api_key=api_keys["Google"])
        # Store the configured genai module itself as the "client"
        clients["Google"] = genai
        client_load_status["Google"] = "✅ Initialized"
    except GoogleAPICoreExceptions.GoogleAPIError as e:
         client_load_status["Google"] = f"❌ Google API Config Error: {e}"
    except Exception as e: # Catch potential auth errors during configure
        client_load_status["Google"] = f"❌ Google Init Error: {e}"
else:
    client_load_status["Google"] = "⚠️ Key Missing"

# Initialize OpenRouter Client (using OpenAI SDK structure)
if api_keys["OpenRouter"]:
    try:
        clients["OpenRouter"] = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_keys["OpenRouter"],
            default_headers={
                 "HTTP-Referer": YOUR_SITE_URL,
                 "X-Title": YOUR_APP_NAME,
            },
        )
        client_load_status["OpenRouter"] = "✅ Initialized"
    except OpenAIError as e:
         client_load_status["OpenRouter"] = f"❌ OpenRouter Config Error: {e}"
    except Exception as e:
         client_load_status["OpenRouter"] = f"❌ Unexpected OpenRouter Init Error: {e}"
else:
    client_load_status["OpenRouter"] = "⚠️ Key Missing"


# --- Model Mapping (Define known models for selection) ---
MODEL_MAPPING = {
    "OpenAI": ["gpt-4.1-mini", "gpt-4.1-nano", "gpt-4.1", "o4-mini", "03"],
    "Anthropic": [
        "claude-3-7-sonnet-20250219",
        "claude-3-5-sonnet-20241022",
        "claude-3-5-haiku-20241022",
        "claude-3-opus-20240229",
    ],
    "Google": [
        "gemini-2.5-pro-preview",
        "models/gemini-1.5-pro-latest",
        "gemini-2.0-flash",
        "models/gemini-1.5-flash-latest",
        "gemini-2.0-flash-lite",
        "gemini-1.5-flash-8b",
        "models/gemini-1.0-pro"
    ],
    "OpenRouter": ["Fetching..."]
}
OPENROUTER_MODELS_CACHE = None

# --- Helper Functions ---

def validate_chunks(chunks, min_length=100):
    """Validate and clean chunks before returning"""
    if not chunks:
        return []
        
    validated = []
    for i, chunk in enumerate(chunks):
        # Skip empty or too-short chunks
        if not chunk or len(chunk.strip()) < min_length:
            logger.warning(f"Skipping chunk {i+1}: too short ({len(chunk.strip()) if chunk else 0} chars)")
            continue
        validated.append(chunk)
    
    # If all chunks were filtered, keep originals as fallback
    if not validated and chunks:
        logger.warning("All chunks were filtered out, keeping originals as fallback")
        return chunks
    
    logger.info(f"Chunk validation: {len(validated)} valid chunks from {len(chunks)} total")
    return validated

def generate_docx_bytes_safe(content_dict):
    """Generate DOCX with comprehensive error handling"""
    try:
        document = docx.Document()
        
        for title, text_content in content_dict.items():
            try:
                # Clean title
                title_str = str(title) if title is not None else "Untitled Section"
                clean_title = ''.join(c for c in title_str if c.isprintable() or c.isspace()).strip()
                if not clean_title:
                    clean_title = "Untitled Section"
                
                # Clean text content
                text_str = str(text_content) if text_content is not None else ""
                
                # More aggressive cleaning for DOCX compatibility
                cleaned_text = ""
                for char in text_str:
                    if ord(char) < 32 and ord(char) not in (9, 10, 13):  # Control chars except tab/newline
                        cleaned_text += ' '
                    elif ord(char) > 127:  # Non-ASCII
                        # Try to keep common unicode chars
                        if ord(char) < 0x10000:  # Basic Multilingual Plane
                            cleaned_text += char
                        else:
                            cleaned_text += '?'  # Replace with placeholder
                    else:
                        cleaned_text += char
                
                # Add to document
                document.add_heading(clean_title, level=1)
                
                # Split very long paragraphs
                if len(cleaned_text) > 10000:
                    # Split into smaller paragraphs
                    chunks = [cleaned_text[i:i+10000] for i in range(0, len(cleaned_text), 10000)]
                    for chunk in chunks:
                        document.add_paragraph(chunk)
                else:
                    document.add_paragraph(cleaned_text)
                    
            except Exception as e:
                logger.error(f"Error adding section '{title}': {e}")
                # Add error notice to document
                document.add_paragraph(f"[Error processing section: {clean_title}]")
        
        # Save to bytes
        bio = io.BytesIO()
        document.save(bio)
        bio.seek(0)
        return bio.getvalue()
        
    except Exception as e:
        logger.error(f"Critical error generating DOCX: {e}")
        # Return None to trigger fallback
        return None

def save_as_text_fallback(content_dict, filepath):
    """Save as plain text if DOCX generation fails"""
    try:
        txt_path = filepath.replace('.docx', '.txt')
        with open(txt_path, 'w', encoding='utf-8') as f:
            for title, content in content_dict.items():
                f.write(f"\n{'='*50}\n{title}\n{'='*50}\n\n")
                f.write(str(content))
                f.write("\n\n")
        st.warning(f"DOCX generation failed. Saved as text file: {txt_path}")
        return txt_path
    except Exception as e:
        st.error(f"Failed to save text fallback: {e}")
        return None

def parse_custom_names(custom_instructions: str) -> Dict[str, str]:
    """Extract correct name spellings from custom instructions"""
    name_mappings = {}
    
    if not custom_instructions:
        return name_mappings
    
    lines = custom_instructions.strip().split('\n')
    for line in lines:
        # Look for patterns like "Name (not OtherName)" or "Correct Name"
        clean_line = line.strip()
        if clean_line and not clean_line.startswith('#'):  # Skip comments
            # Extract the correct name (before any parentheses)
            name = clean_line.split('(')[0].strip()
            if name:
                # Create variations for matching
                name_lower = name.lower()
                name_mappings[name_lower] = name
                # Also add without spaces for matching
                name_no_space = name.replace(' ', '').lower()
                if name_no_space != name_lower:
                    name_mappings[name_no_space] = name
    
    logger.info(f"Parsed {len(name_mappings)} custom name mappings from instructions")
    return name_mappings

def estimate_llm_cost(text: str, provider: str, model: str) -> Dict[str, float]:
    """Estimate cost for LLM processing"""
    # Rough token estimation (1 token ≈ 4 characters)
    estimated_tokens = len(text) / 4
    
    # Cost per 1K tokens (approximate, update with actual costs)
    cost_per_1k = {
        ("OpenAI", "gpt-4"): {"input": 0.03, "output": 0.06},
        ("OpenAI", "gpt-4-turbo"): {"input": 0.01, "output": 0.03},
        ("OpenAI", "gpt-3.5-turbo"): {"input": 0.0005, "output": 0.0015},
        ("Anthropic", "claude-3-opus"): {"input": 0.015, "output": 0.075},
        ("Anthropic", "claude-3-sonnet"): {"input": 0.003, "output": 0.015},
        ("Anthropic", "claude-3-haiku"): {"input": 0.00025, "output": 0.00125},
        ("Google", "gemini-pro"): {"input": 0.00025, "output": 0.0005},
    }
    
    # Find matching cost or use default
    model_base = model.split('-')[0] if '-' in model else model
    provider_model = (provider, model_base)
    costs = cost_per_1k.get(provider_model, {"input": 0.001, "output": 0.002})
    
    # Estimate total cost (assume output is 20% of input)
    input_cost = (estimated_tokens / 1000) * costs["input"]
    output_cost = (estimated_tokens * 0.2 / 1000) * costs["output"]
    
    return {
        "estimated_tokens": int(estimated_tokens),
        "estimated_cost": round(input_cost + output_cost, 4),
        "breakdown": f"Input: ${input_cost:.4f}, Output: ${output_cost:.4f}"
    }

@st.cache_data(ttl=3600)
def fetch_openrouter_models(api_key):
    global OPENROUTER_MODELS_CACHE
    if OPENROUTER_MODELS_CACHE is not None and (not OPENROUTER_MODELS_CACHE or not OPENROUTER_MODELS_CACHE[0].startswith("Error:")):
        return OPENROUTER_MODELS_CACHE
    if not api_key:
        return ["Error: OpenRouter API Key missing"]
    st.info("Fetching OpenRouter models list...")
    try:
        headers = {"Authorization": f"Bearer {api_key}"}
        response = requests.get("https://openrouter.ai/api/v1/models", headers=headers, timeout=20)
        response.raise_for_status()
        models_data = response.json().get("data", [])
        model_ids = sorted([model.get("id") for model in models_data if model.get("id")])
        if not model_ids: result = ["Error: No models found in OpenRouter response"]
        else: result = model_ids
        st.info(f"OpenRouter models list fetched ({len(result)} models).")
        OPENROUTER_MODELS_CACHE = result
        return result
    except requests.exceptions.Timeout:
        st.error("Failed to fetch OpenRouter models: Request timed out.")
        result = ["Error: Request Timed Out"]
        OPENROUTER_MODELS_CACHE = result
        return result
    except requests.exceptions.RequestException as e:
        st.error(f"Failed to fetch OpenRouter models: Network error - {e}")
        result = [f"Error: Connection Failed"]
        OPENROUTER_MODELS_CACHE = result
        return result
    except Exception as e:
        st.error(f"Error processing OpenRouter models response: {e}")
        result = [f"Error: Processing Failed"]
        OPENROUTER_MODELS_CACHE = result
        return result

def load_instruction(file_path):
    try:
        normalized_path = os.path.normpath(file_path)
        with open(normalized_path, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        st.warning(f"Instruction file not found: {normalized_path}")
        return ""
    except Exception as e:
        st.error(f"Error reading instruction file {normalized_path}: {e}")
        return ""

def call_llm_with_retry(prompt, provider, model_name, max_retries=3):
    """Call LLM with retry logic for transient errors"""
    retry_delays = [1, 2, 4]  # Exponential backoff
    last_error = None
    
    for attempt in range(max_retries):
        try:
            result = call_llm(prompt, provider, model_name)
            
            # Check if it's a retryable error
            if result.startswith("Error:"):
                error_lower = result.lower()
                # Don't retry on permanent errors
                if any(x in error_lower for x in ["invalid", "authentication", "not found"]):
                    logger.error(f"Permanent error in LLM call: {result}")
                    return result
                
                # Retry on transient errors
                if any(x in error_lower for x in ["rate", "network", "timeout", "connection"]):
                    if attempt < max_retries - 1:
                        delay = retry_delays[attempt]
                        st.warning(f"Transient error, retrying in {delay}s... ({attempt+1}/{max_retries})")
                        logger.warning(f"Retrying LLM call after error: {result}")
                        time.sleep(delay)
                        continue
                    
            logger.info(f"Successful LLM call to {provider}/{model_name}")
            return result
            
        except Exception as e:
            last_error = str(e)
            logger.error(f"Exception in LLM call attempt {attempt+1}: {e}")
            if attempt < max_retries - 1:
                delay = retry_delays[attempt]
                st.warning(f"Error: {e}. Retrying in {delay}s... ({attempt+1}/{max_retries})")
                time.sleep(delay)
            else:
                logger.error(f"All LLM retry attempts failed. Last error: {last_error}")
                return f"Error after {max_retries} attempts: {last_error}"
    
    return f"Error: Max retries exceeded"

def call_llm(prompt, provider, model_name):
    st.write(f"Attempting LLM call: Provider='{provider}', Model='{model_name}' (Using Default Temperature)")
    if not provider or not model_name:
        st.error("LLM Call Error: Provider or model name missing.")
        return "Error: Provider or model not specified."
    client = clients.get(provider)
    if not client:
        st.error(f"LLM Call Error: Client for '{provider}' not available. Check API key and initialization status in the sidebar.")
        return f"Error: Client for {provider} not initialized."
    if model_name.startswith("Error:") or model_name == "Fetching...":
         st.error(f"LLM Call Error: Invalid model selected for {provider}: {model_name}")
         return f"Error: Invalid model '{model_name}' selected."
    try:
        if provider == "OpenAI" or provider == "OpenRouter":
            response = client.chat.completions.create(
                model=model_name,
                messages=[{"role": "user", "content": prompt}]
            )
            result = response.choices[0].message.content
            st.write(f"LLM call successful ({provider})")
            return result.strip() if result else ""
        elif provider == "Anthropic":
            system_prompt = ""
            user_prompt = prompt
            prompt_lines = prompt.split('\n', 1)
            if len(prompt_lines) > 1 and len(prompt_lines[0]) < 300 and \
               any(keyword in prompt_lines[0].lower() for keyword in ["instructions", "role", "task", "goal"]):
                system_prompt = prompt_lines[0]
                user_prompt = prompt_lines[1]
            max_output_tokens = 4096
            if "claude-3-5" in model_name or "claude-3-7" in model_name:
                max_output_tokens = 8192
            response = client.messages.create(
                model=model_name,
                max_tokens=max_output_tokens,
                system=system_prompt if system_prompt else None,
                messages=[{"role": "user", "content": user_prompt}]
            )
            if response.stop_reason == "max_tokens":
                 st.warning(f"Anthropic response stopped due to max_tokens ({max_output_tokens}). Output might be truncated.")
            result_text = ""
            if isinstance(response.content, list):
                 result_text = "".join([block.text for block in response.content if hasattr(block, 'text')])
            elif hasattr(response.content, 'text'):
                 result_text = response.content.text
            else:
                 st.warning("Unexpected Anthropic response content structure.")
                 result_text = str(response.content)
            st.write(f"LLM call successful ({provider})")
            return result_text.strip() if result_text else ""
        elif provider == "Google":
            safety_settings = {
                HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
            }
            model = client.GenerativeModel(
                model_name=model_name,
                safety_settings=safety_settings
            )
            response = model.generate_content(prompt)
            prompt_feedback = getattr(response, 'prompt_feedback', None)
            block_reason = getattr(prompt_feedback, 'block_reason', None) if prompt_feedback else None
            if block_reason:
                 safety_ratings = getattr(prompt_feedback, 'safety_ratings', [])
                 st.error(f"Google API call blocked. Reason: {block_reason}. Safety Ratings: {safety_ratings}")
                 try: st.warning(f"Full Google Response (Debug): {response}")
                 except Exception: st.warning("Could not display full Google response object.")
                 return f"Error: Google API Blocked Prompt ({block_reason})"
            try:
                result = response.text
            except ValueError as ve:
                 candidates = getattr(response, 'candidates', [])
                 if candidates:
                     st.warning(f"Google response generated candidate(s) but accessing '.text' failed ({ve}). Candidates: {candidates}")
                     if candidates[0].content and candidates[0].content.parts:
                         result = "".join([part.text for part in candidates[0].content.parts if hasattr(part, 'text')])
                         if result: return result.strip()
                 st.error(f"Google API call failed: Could not extract text content. Error: {ve}")
                 return f"Error: Google response missing text content ({ve})."
            except Exception as e_text:
                 st.error(f"Unexpected error accessing Google response text: {e_text}")
                 return f"Error: Failed to access Google response text ({e_text})."
            st.write(f"LLM call successful ({provider})")
            return result.strip() if result else ""
        else:
            st.error(f"LLM Call Error: Provider '{provider}' logic not implemented.")
            return f"Error: Unknown provider {provider}."
    except OpenAIError as e:
        st.error(f"API Error ({provider} via OpenAI SDK): {e}")
        error_detail = str(e).lower()
        if "context_length_exceeded" in error_detail: return f"Error: Context length exceeded for {model_name}."
        if "model_not_found" in error_detail: return f"Error: Model not found: {model_name}."
        if "authentication" in error_detail: return f"Error: Authentication failed for {provider}. Check API Key."
        return f"Error: OpenAI/OR Error - {e}"
    except AnthropicAPIError as e:
        st.error(f"API Error (Anthropic): {e}")
        error_detail = str(e).lower()
        if "authentication" in error_detail: return "Error: Authentication failed for Anthropic. Check API Key."
        if "quota" in error_detail or "rate_limit" in error_detail: return f"Error: Anthropic Quota/Rate Limit Exceeded - {e}"
        if "invalid_request" in error_detail: return f"Error: Invalid Request for Anthropic model {model_name} - {e}"
        return f"Error: Anthropic Error - {e}"
    except GoogleAPICoreExceptions.PermissionDenied as e:
         st.error(f"API Error (Google): Permission Denied. Check API Key & API enabled status. Details: {e}")
         return f"Error: Google Permission Denied."
    except GoogleAPICoreExceptions.ResourceExhausted as e:
        st.error(f"API Error (Google): Resource Exhausted (Quota Limit?). Details: {e}")
        return f"Error: Google Quota Exceeded."
    except GoogleAPICoreExceptions.InvalidArgument as e:
         st.error(f"API Error (Google): Invalid Argument. Check model name ('{model_name}')/prompt. Details: {e}")
         return f"Error: Google Invalid Argument."
    except GoogleAPICoreExceptions.GoogleAPIError as e:
        st.error(f"API Error (Google): {e}")
        return f"Error: Google API Error - {e}"
    except requests.exceptions.RequestException as e:
         st.error(f"Network Error for {provider}: {e}")
         return f"Error: Network Connection Failed ({provider})."
    except Exception as e:
        st.error(f"An unexpected error occurred during the LLM call ({provider} - {model_name}): {e}")
        st.error(traceback.format_exc())
        return f"Unexpected Error ({provider}) - {e}"

def parse_uploaded_file(uploaded_file):
    if uploaded_file is None:
        return None, "No file uploaded."
    file_name = uploaded_file.name
    file_type = ""
    if '.' in file_name:
        file_type = file_name.split('.')[-1].lower()
    st.info(f"Parsing file: '{file_name}' (Detected type: '{file_type}')")
    try:
        bytes_data = uploaded_file.getvalue()
        if not bytes_data:
             st.warning(f"Uploaded file '{file_name}' appears to be empty.")
             return "", None
             
        # Check file size
        file_size = len(bytes_data)
        file_size_mb = file_size / (1024 * 1024)
        MAX_FILE_SIZE_MB = 50
        
        logger.info(f"Processing file: {file_name}, Size: {file_size_mb:.1f}MB")
        
        if file_size_mb > MAX_FILE_SIZE_MB:
            error_msg = f"File too large: {file_size_mb:.1f}MB. Maximum allowed: {MAX_FILE_SIZE_MB}MB"
            logger.error(error_msg)
            return None, error_msg
        
        st.info(f"File size: {file_size_mb:.1f}MB")
        if file_type == 'pdf':
            text = ""
            is_ocr_needed = False
            try:
                with fitz.open(stream=bytes_data, filetype="pdf") as doc:
                    if not doc.page_count:
                         st.warning(f"PDF file '{file_name}' has zero pages.")
                         return "", None
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        page_text_sorted = page.get_text("text", sort=True).strip()
                        page_text_simple = page.get_text("simple").strip()
                        page_text = page_text_sorted if page_text_sorted else page_text_simple
                        if not page_text and page.get_images(full=True):
                             is_ocr_needed = True
                        if page_text:
                             text += page_text + "\n"
                        if page_num < len(doc) - 1:
                              text += "\n--- Page Break ---\n"
            except fitz.fitz.FileDataError as fe:
                 return None, f"Error parsing PDF '{file_name}': File may be corrupt, password protected, or not a valid PDF format. Details: {fe}"
            except Exception as fitz_e:
                 return None, f"Unexpected error during PDF processing: {fitz_e}\n{traceback.format_exc()}"
            meaningful_text_check = text.replace("--- Page Break ---", "").strip()
            if not meaningful_text_check:
                 if is_ocr_needed:
                      return None, f"Failed to extract text from PDF: '{file_name}'. The document appears to be image-based and requires OCR (Optical Character Recognition), which this tool does not perform."
                 else:
                      st.warning(f"Parsed PDF '{file_name}', but no text content was extracted. The file might be empty or have an unusual structure.")
                      return "", None
            st.success(f"Successfully parsed PDF: '{file_name}'")
            return text, None
        elif file_type == 'docx':
            try:
                doc_io = io.BytesIO(bytes_data)
                doc = docx.Document(doc_io)
                paragraphs_text = [para.text for para in doc.paragraphs if para.text and para.text.strip()]
                text = "\n".join(paragraphs_text)
                if not text.strip():
                     st.warning(f"DOCX file '{file_name}' parsed, but contained no text in its paragraphs.")
                     return "", None
                st.success(f"Successfully parsed DOCX: '{file_name}'")
                return text, None
            except Exception as docx_e:
                 return None, f"Error opening or parsing DOCX file '{file_name}'. It might be corrupted or not a valid DOCX format. Error: {docx_e}\n{traceback.format_exc()}"
        elif file_type == 'txt':
            try:
                text = bytes_data.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    text = bytes_data.decode("cp1252")
                    st.warning("Decoded TXT file using CP1252 encoding (UTF-8 failed). Check for garbled characters.")
                except Exception as e_decode:
                   st.error(f"Failed to decode TXT file '{file_name}' with UTF-8 or CP1252: {e_decode}")
                   return None, f"TXT decode error: Could not decode file content ({e_decode})."
            if not text.strip():
                 st.warning(f"TXT file '{file_name}' parsed, but was empty or contained only whitespace.")
                 return "", None
            st.success(f"Successfully parsed TXT: '{file_name}'")
            return text, None
        else:
            error_msg = f"Unsupported file type: '.{file_type}'. Please upload a file with a .txt, .pdf, or .docx extension."
            st.error(error_msg)
            return None, error_msg
    except Exception as e:
        error_msg = f"An unexpected error occurred during file parsing for '{file_name}': {e}\n{traceback.format_exc()}"
        st.error(error_msg)
        return None, error_msg

# Note: TopicBasedChunker functionality is now provided by EnhancedTopicBasedChunker
# imported from enhanced_topic_chunker_fixed.py 
def enhanced_chunk_text(text: str, 
                       provider: str = None, 
                       model: str = None,
                       chunk_size: int = 8000, 
                       chunk_overlap: int = 400) -> List[str]:
    """
    Enhanced chunking function that replaces the original chunk_text.
    Falls back to original method if topic-based chunking fails.
    """
    
    # If LLM details not provided, or text is very short, fall back to original method
    if not provider or not model:
        st.warning("No LLM specified for topic chunking, using default (RecursiveCharacterTextSplitter) method...")
        return chunk_text(text, chunk_size, chunk_overlap) # Call original chunk_text
    if not text or len(text.strip()) < 500: # Arbitrary short text threshold
        st.info("Text is very short, using default (RecursiveCharacterTextSplitter) method for chunking.")
        return chunk_text(text, chunk_size, chunk_overlap)

    st.info(f"Attempting Enhanced Topic-Based Chunking with {provider}/{model}...")
    try:
        # Use enhanced topic-based chunking with quality guarantees
        chunker = EnhancedTopicBasedChunker(
            llm_provider=provider,
            llm_model=model,
            target_chunk_size=chunk_size,
            min_chunk_size=2000,
            max_chunk_size=12000,
            overlap_size=chunk_overlap,
            quality_threshold=0.7,
            call_llm_func=call_llm_with_retry  # Pass the LLM function directly
        )
        
        chunks = chunker.chunk_transcript(text)
        
        # Validate chunks
        if chunks and len(chunks) > 0:
            chunk_sizes = [len(chunk) for chunk in chunks]
            avg_size_val = int(sum(chunk_sizes) / len(chunks)) if len(chunks) > 0 else 0
            min_size_val = min(chunk_sizes) if chunk_sizes else 0
            max_size_val = max(chunk_sizes) if chunk_sizes else 0

            st.success(f"Enhanced topic-based chunking complete: {len(chunks)} chunks created.")
            st.write(f"Chunk stats - Avg size: {avg_size_val} chars, Min: {min_size_val}, Max: {max_size_val}")
            
            # Display processing stats from the enhanced chunker
            stats = chunker.get_stats()
            st.info(f"Processing stats - Time: {stats['processing_time']:.1f}s, "
                   f"LLM calls: {stats['llm_calls']}, "
                   f"Fallbacks: {stats['fallback_count']}")
            
            # Validate chunks before returning
            validated_chunks = validate_chunks(chunks, min_length=100)
            logger.info(f"Enhanced topic-based chunking completed: {len(validated_chunks)} validated chunks")
            return validated_chunks
        else:
            st.warning("Topic chunking produced no chunks or failed. Falling back to default method.")
            logger.warning("Topic chunking failed, using fallback method")
            return chunk_text(text, chunk_size, chunk_overlap) # Fallback
            
    except Exception as e:
        st.error(f"Topic-based chunking failed with an unexpected error: {e}")
        st.error(traceback.format_exc())
        st.warning("Falling back to default (RecursiveCharacterTextSplitter) chunking method...")
        return chunk_text(text, chunk_size, chunk_overlap) # Fallback


def chunk_text(text, chunk_size=8000, chunk_overlap=400):
    """Splits the text into overlapping chunks for LLM processing."""
    st.info(f"Chunking text using RecursiveCharacterTextSplitter (Size: {chunk_size}, Overlap: {chunk_overlap})...")
    if not text or text.isspace():
        st.warning("Input text for chunking is empty or only whitespace.")
        return []
    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            is_separator_regex=False,
            separators=["\n\n", "\n", ". ", "! ", "? ", "; ", ": ", ",", " ", ""],
            keep_separator=True
        )
        chunks = text_splitter.split_text(text)
        chunks = [chunk for chunk in chunks if chunk and not chunk.isspace()]
        if not chunks:
            st.warning("Text chunking resulted in zero non-whitespace chunks.")
            logger.warning("Default chunking produced no valid chunks")
            if text and not text.isspace():
                st.info("Returning original text as a single chunk.")
                return [text]
            else:
                return []
        
        # Validate chunks
        validated_chunks = validate_chunks(chunks, min_length=50)
        st.write(f"Text divided into {len(validated_chunks)} validated chunks (RecursiveCharacterTextSplitter).")
        logger.info(f"Default chunking completed: {len(validated_chunks)} chunks")
        return validated_chunks
    except Exception as e:
        st.error(f"An error occurred during text chunking: {e}")
        st.error(traceback.format_exc())
        return []

def save_bytes_to_file(file_bytes, full_path):
    if file_bytes is None:
        st.error(f"Attempted to save 'None' to {os.path.basename(full_path)}. Document generation likely failed.")
        return None
    try:
        dir_name = os.path.dirname(full_path)
        if dir_name:
             os.makedirs(dir_name, exist_ok=True)
        base, ext = os.path.splitext(full_path)
        counter = 1
        final_path = full_path
        while os.path.exists(final_path):
            final_path = f"{base}_({counter}){ext}"
            counter += 1
        with open(final_path, 'wb') as f:
            f.write(file_bytes)
        st.success(f"✔️ File successfully saved to: {final_path}")
        return final_path
    except PermissionError:
        st.error(f"Permission denied: Cannot write to the directory '{os.path.dirname(full_path)}'. Please check folder permissions.")
        return None
    except OSError as oe:
        st.error(f"Failed to save file to {full_path} (or variant): OS Error - {oe}")
        st.error(traceback.format_exc())
        return None
    except Exception as e:
        st.error(f"An unexpected error occurred while saving file to {full_path} (or variant): {e}")
        st.error(traceback.format_exc())
        return None

# --- Load Instructions ---
INSTRUCTIONS_DIR = "instructions"
global_context_instructions_path = os.path.join(INSTRUCTIONS_DIR, "global_context_instructions.txt")
standard_summary_instructions_path = os.path.join(INSTRUCTIONS_DIR, "standard_summary_instructions.txt")
global_context_instructions = load_instruction(global_context_instructions_path)
standard_instructions = load_instruction(standard_summary_instructions_path)
role_files = {
    "Plaintiff": os.path.join(INSTRUCTIONS_DIR, "plaintiff_outline.txt"),
    "Defendant": os.path.join(INSTRUCTIONS_DIR, "defendant_outline.txt"),
    "Expert Witness": os.path.join(INSTRUCTIONS_DIR, "expert_witness_outline.txt"),
    "Fact Witness": os.path.join(INSTRUCTIONS_DIR, "fact_witness_outline.txt"),
}
instruction_load_status = {}
instruction_load_status["Global Context"] = bool(global_context_instructions)
instruction_load_status["Standard Summary"] = bool(standard_instructions)
for role, path in role_files.items():
    instruction_load_status[f"{role} Outline"] = os.path.exists(path)

# --- Sidebar Setup ---
st.sidebar.title("⚙️ Config & Status")
st.sidebar.markdown("---")
available_providers = [p for p, status in client_load_status.items() if "✅" in status]

if not available_providers:
      st.sidebar.error("No LLM providers initialized successfully. Check API keys in the `.env` file.")
      selected_provider_context = None
      selected_model_context = None
      selected_provider_final = None
      selected_model_final = None
else:
    st.sidebar.subheader("1. Context & Segment LLM")
    st.sidebar.caption("Used for initial context generation and summarizing text chunks. Can often be a faster/cheaper model.")
    
    # CHANGE 2: Set default to OpenRouter if available
    default_provider_index = 0
    if "OpenRouter" in available_providers:
        default_provider_index = available_providers.index("OpenRouter")
    
    selected_provider_context = st.sidebar.selectbox(
        "Select Provider (Context/Segments):", options=available_providers, index=default_provider_index, key="provider_selector_context",
        help="Choose the LLM service for context generation and segment summaries."
    )
    models_for_provider_context = []
    if selected_provider_context:
        if selected_provider_context == "OpenRouter":
            openrouter_key = api_keys.get("OpenRouter")
            models_for_provider_context = fetch_openrouter_models(openrouter_key) if openrouter_key else ["Error: OpenRouter Key Missing"]
        else:
            models_for_provider_context = MODEL_MAPPING.get(selected_provider_context, ["Error: No models defined"])
        if not isinstance(models_for_provider_context, list): models_for_provider_context = ["Error: Invalid model list format"]
        if models_for_provider_context:
            default_index_context = 0
            if selected_provider_context == "Google" and "gemini-2.5-pro-preview" in models_for_provider_context:
                default_index_context = models_for_provider_context.index("gemini-2.5-pro-preview")
            elif selected_provider_context == "OpenAI":
                # Try to find a suitable default model for context processing
                if "gpt-4.1-mini" in models_for_provider_context: 
                    default_index_context = models_for_provider_context.index("gpt-4.1-mini")
                elif "o4-mini" in models_for_provider_context: 
                    default_index_context = models_for_provider_context.index("o4-mini")
                elif "gpt-4.1-nano" in models_for_provider_context:
                    default_index_context = models_for_provider_context.index("gpt-4.1-nano")
            elif selected_provider_context == "Anthropic" and "claude-3-5-haiku-20241022" in models_for_provider_context:
                default_index_context = models_for_provider_context.index("claude-3-5-haiku-20241022")
            # CHANGE 2: Set OpenRouter default to google/gemini-2.5-pro-preview
            elif selected_provider_context == "OpenRouter" and not models_for_provider_context[0].startswith("Error"):
                if "google/gemini-2.5-pro-preview" in models_for_provider_context:
                    default_index_context = models_for_provider_context.index("google/gemini-2.5-pro-preview")
                else:
                    # Fallback to other cheap models if the specific one isn't available
                    cheap_or_models = ['google/gemini-flash-1.5', 'anthropic/claude-3-haiku-20240722', 'openai/gpt-3.5-turbo', 'mistralai/mistral-7b-instruct']
                    for m_or in cheap_or_models:
                         if m_or in models_for_provider_context: default_index_context = models_for_provider_context.index(m_or); break
            selected_model_context = st.sidebar.selectbox(
                 "Select Model (Context/Segments):", options=models_for_provider_context, index=default_index_context,
                 key=f"model_selector_context_{selected_provider_context}",
                 help=f"Choose the specific model from {selected_provider_context} for context generation and segment summarization."
            )
        else:
            st.sidebar.warning(f"No models available or loaded for {selected_provider_context}.")
            selected_model_context = None
    else:
        selected_model_context = None

    st.sidebar.markdown("---")
    st.sidebar.subheader("2. Final Synthesis LLM")
    st.sidebar.caption("Used for combining segment summaries into the final structured output. Often benefits from a more capable model.")
    selected_provider_final = st.sidebar.selectbox(
        "Select Provider (Final Synthesis):", options=available_providers, index=0, key="provider_selector_final",
        help="Choose the LLM service for the final summary synthesis step."
    )
    models_for_provider_final = []
    if selected_provider_final:
        if selected_provider_final == "OpenRouter":
            openrouter_key = api_keys.get("OpenRouter")
            models_for_provider_final = fetch_openrouter_models(openrouter_key) if openrouter_key else ["Error: OpenRouter Key Missing"]
        else:
            models_for_provider_final = MODEL_MAPPING.get(selected_provider_final, ["Error: No models defined"])
        if not isinstance(models_for_provider_final, list): models_for_provider_final = ["Error: Invalid model list format"]
        if models_for_provider_final:
            default_index_final = 0
            if selected_provider_final == "Google" and "gemini-2.5-pro-preview" in models_for_provider_final:
                default_index_final = models_for_provider_final.index("gemini-2.5-pro-preview")
            elif selected_provider_final == "OpenAI":
                # Try to find a suitable default model for final synthesis
                if "gpt-4.1" in models_for_provider_final:
                    default_index_final = models_for_provider_final.index("gpt-4.1")
                elif "03" in models_for_provider_final:
                    default_index_final = models_for_provider_final.index("03")
                elif "gpt-4.1-mini" in models_for_provider_final:
                    default_index_final = models_for_provider_final.index("gpt-4.1-mini")
            elif selected_provider_final == "Anthropic" and "claude-3-7-sonnet-20250219" in models_for_provider_final:
                default_index_final = models_for_provider_final.index("claude-3-7-sonnet-20250219")
            elif selected_provider_final == "OpenRouter" and not models_for_provider_final[0].startswith("Error"):
                capable_or_models = ['openai/gpt-4o', 'anthropic/claude-3.7-sonnet', 'google/gemini-pro-1.5', 'mistralai/mixtral-8x22b-instruct']
                for m_or_cap in capable_or_models: # Renamed m to m_or_cap
                     if m_or_cap in models_for_provider_final: default_index_final = models_for_provider_final.index(m_or_cap); break
            selected_model_final = st.sidebar.selectbox(
                 "Select Model (Final Synthesis):", options=models_for_provider_final, index=default_index_final,
                 key=f"model_selector_final_{selected_provider_final}",
                 help=f"Choose the specific model from {selected_provider_final} for the final synthesis step."
            )
        else:
            st.sidebar.warning(f"No models available or loaded for {selected_provider_final}.")
            selected_model_final = None
    else:
        selected_model_final = None

st.sidebar.markdown("---")
st.sidebar.subheader("Selected LLMs (Using Default Temperature):")
if available_providers:
    if selected_provider_context and selected_model_context and not selected_model_context.startswith("Error:"):
        st.sidebar.write(f"Context/Segments: `{selected_provider_context} / {selected_model_context}`")
    elif selected_provider_context:
        st.sidebar.warning(f"Context/Segments: Invalid model selected for `{selected_provider_context}`.")
    else:
         st.sidebar.warning("Context/Segment LLM not selected.")
    if selected_provider_final and selected_model_final and not selected_model_final.startswith("Error:"):
        st.sidebar.write(f"Final Synthesis: `{selected_provider_final} / {selected_model_final}`")
    elif selected_provider_final:
        st.sidebar.warning(f"Final Synthesis: Invalid model selected for `{selected_provider_final}`.")
    else:
         st.sidebar.warning("Final Synthesis LLM not selected.")

st.sidebar.markdown("---")
st.sidebar.subheader("API Key & Client Status:")
for provider, status in client_load_status.items():
    if "✅" in status: st.sidebar.success(f"{provider}: {status}")
    elif "⚠️" in status: st.sidebar.warning(f"{provider}: {status}")
    else: st.sidebar.error(f"{provider}: {status}")

st.sidebar.markdown("---")
st.sidebar.subheader("Instruction Files Status:")
if instruction_load_status["Global Context"]: st.sidebar.success("✅ Global context instructions loaded.")
else: st.sidebar.warning("⚠️ Global context instructions missing or empty.")
if instruction_load_status["Standard Summary"]: st.sidebar.success("✅ Standard summary instructions loaded.")
else: st.sidebar.warning("⚠️ Standard summary instructions missing or empty.")
all_roles_loaded = True
for role, path in role_files.items():
    role_key = f"{role} Outline"
    exists = instruction_load_status.get(role_key, False)
    if exists:
        st.sidebar.success(f"✅ {role} outline file found.")
    else:
        st.sidebar.error(f"❌ {role} outline file MISSING ({os.path.basename(path)})")
        all_roles_loaded = False
if not all_roles_loaded:
    st.sidebar.error("One or more required role outline files are missing from the 'instructions' folder.")

# Initialize session state for progress tracking
if 'processing_state' not in st.session_state:
    st.session_state.processing_state = {
        'last_chunk_processed': 0,
        'partial_summaries': [],
        'global_context': None,
        'transcript_text': None,
        'processing_start_time': None
    }

# --- Streamlit UI (Main Area) ---
st.title("📄 Deposition Transcript Summarizer")
st.markdown(f"""
Upload a transcript (.txt, .pdf, .docx), select the deponent's role, configure **two LLMs** in the sidebar (one LLM for context/segments, one for final synthesis), add optional instructions (especially correct name spellings), and click generate. Models will use their default temperature settings.

**Output Location:** `{OUTPUT_DIRECTORY}` (Ensure this directory exists and is writable). Files will not be overwritten; a number like `_(1)` will be appended if a file with the same name exists.
""")

col1, col2 = st.columns([2, 1])
with col1:
    uploaded_file = st.file_uploader(
        "1. Upload Transcript", type=["txt", "pdf", "docx"], key="file_uploader"
    )
    deponent_role = st.selectbox(
        "2. Select Deponent Role:", options=list(role_files.keys()), key="role_selector"
    )
with col2:
    custom_instructions = st.text_area(
        "3. Optional: Add Custom Instructions (Focus on Names)", height=180,
        placeholder="Enter correctly spelled names (people, places, companies). The AI *must* use these exact spellings in the output.\n\nExample:\nJohnathan P. Smyth (not Jon Smith)\nACME Corporation Ltd.\n123 Oak Street, Anytown",
        key="custom_instructions",
        help="Provide a list of key names with their correct spellings. The LLMs will be explicitly instructed to use these exact spellings for any corresponding names found in the transcript. This helps ensure consistency and accuracy in the summaries. You can also add other general instructions here."
    )

st.markdown("---")
process_button_disabled = (
    uploaded_file is None or not available_providers or
    selected_provider_context is None or selected_model_context is None or
    selected_model_context.startswith("Error:") or selected_model_context == "Fetching..." or
    selected_provider_final is None or selected_model_final is None or
    selected_model_final.startswith("Error:") or selected_model_final == "Fetching..."
)
process_button_tooltip = "Please upload a transcript file and select valid LLM providers/models for both stages in the sidebar to enable generation." if process_button_disabled else "Start processing the transcript and generate summaries."
process_button = st.button(
    "🚀 Generate & Save Summaries", type="primary", disabled=process_button_disabled,
    key="generate_button", help=process_button_tooltip
)

if process_button:
    # Show cost estimation
    if uploaded_file and selected_provider_context and selected_model_context:
        # Get file content for estimation
        temp_text, temp_error = parse_uploaded_file(uploaded_file)
        if temp_text and not temp_error:
            cost_estimate = estimate_llm_cost(temp_text, selected_provider_context, selected_model_context)
            with st.expander("💰 Cost Estimation"):
                st.write(f"**Estimated tokens:** {cost_estimate['estimated_tokens']:,}")
                st.write(f"**Estimated cost:** ${cost_estimate['estimated_cost']}")
                st.write(f"**Breakdown:** {cost_estimate['breakdown']}")
                st.info("This is a rough estimate. Actual costs may vary.")
    
    # Parse custom names for validation
    custom_name_mappings = parse_custom_names(custom_instructions)
    if custom_name_mappings:
        st.info(f"Parsed {len(custom_name_mappings)} custom name spellings for validation")
    
    # Store processing start time
    st.session_state.processing_state['processing_start_time'] = datetime.now()
    logger.info(f"Starting processing at {st.session_state.processing_state['processing_start_time']}")
    
    st.info(f"Starting Process...")
    if not available_providers: st.error("Critical Error: No LLM providers configured."); st.stop()
    if not selected_provider_context or not selected_model_context or selected_model_context.startswith("Error:") or selected_model_context == "Fetching...": st.error("Pre-check Failed: Invalid Context/Segment LLM selection."); st.stop()
    if not selected_provider_final or not selected_model_final or selected_model_final.startswith("Error:") or selected_model_final == "Fetching...": st.error("Pre-check Failed: Invalid Final Synthesis LLM selection."); st.stop()
    if selected_provider_context not in clients or not client_load_status.get(selected_provider_context, "").startswith("✅"): st.error(f"Client initialization error for Context/Segments provider '{selected_provider_context}'. Check API key/status."); st.stop()
    if selected_provider_final not in clients or not client_load_status.get(selected_provider_final, "").startswith("✅"): st.error(f"Client initialization error for Final Synthesis provider '{selected_provider_final}'. Check API key/status."); st.stop()
    if not standard_instructions: st.error("Pre-check Failed: Standard summary instructions missing or empty."); st.stop()
    if not global_context_instructions: st.error("Pre-check Failed: Global context instructions missing or empty."); st.stop()
    selected_role_file = role_files.get(deponent_role)
    if not selected_role_file or not os.path.exists(selected_role_file): st.error(f"Pre-check Failed: Outline file for role '{deponent_role}' is missing."); st.stop()
    role_outline_instructions = load_instruction(selected_role_file)
    if not role_outline_instructions: st.error(f"Pre-check Failed: Outline file for '{deponent_role}' could not be loaded or is empty."); st.stop()
    try:
        os.makedirs(OUTPUT_DIRECTORY, exist_ok=True)
        test_file_path = os.path.join(OUTPUT_DIRECTORY, ".streamlit_permission_test")
        with open(test_file_path, "w") as f_test: f_test.write("test")
        os.remove(test_file_path)
        st.info(f"Output directory checked and seems writable: {OUTPUT_DIRECTORY}")
    except PermissionError:
         st.error(f"Pre-check Failed: Permission denied writing to the output directory: {OUTPUT_DIRECTORY}. Please check folder permissions.")
         st.stop()
    except Exception as e:
        st.error(f"Pre-check Failed: Could not create, access, or write to output directory: {OUTPUT_DIRECTORY}. Error: {e}")
        st.stop()

    st.markdown("---")
    st.subheader("📊 Processing Status:")
    status_placeholder = st.empty()
    progress_bar = st.progress(0.0)
    saved_path1 = None
    saved_path2 = None
    save_success1_flag = False
    save_success2_flag = False

    def update_status(percentage, text_msg): # Renamed text to text_msg
        clamped_percentage_float = max(0.0, min(1.0, percentage / 100.0))
        text_str = str(text_msg) if text_msg is not None else ""
        progress_bar.progress(clamped_percentage_float, text=text_str)
        text_lower = text_str.lower()
        if "error" in text_lower or "failed" in text_lower or "❌" in text_str: status_placeholder.error(text_str)
        elif "warning" in text_lower or "issue" in text_lower or "⚠️" in text_str: status_placeholder.warning(text_str)
        elif "✔️" in text_str: status_placeholder.success(text_str)
        else: status_placeholder.info(text_str)

    try:
        update_status(0, f"Starting process... Ctx/Seg Model: {selected_model_context}, Final Model: {selected_model_final}")
        update_status(5, f"Step 1/7: Parsing '{uploaded_file.name}'...")
        transcript_text, error = parse_uploaded_file(uploaded_file)
        if error:
            update_status(5, f"❌ Step 1 Failed: File Parsing Error - {error}")
            st.stop()
        if transcript_text is None or not transcript_text.strip():
             update_status(5, f"❌ Step 1 Failed: Parsing '{uploaded_file.name}' resulted in empty text content. Cannot proceed.")
             st.stop()
        update_status(10, "✔️ Step 1: File parsed successfully.")

        update_status(15, f"Step 2/7: Generating global context ({selected_model_context})...")
        context_prompt = f"""BEGIN GLOBAL CONTEXT INSTRUCTIONS:
---
{global_context_instructions}
---
END GLOBAL CONTEXT INSTRUCTIONS.

BEGIN USER'S CUSTOM INSTRUCTIONS (Pay close attention to required name spellings):
---
{custom_instructions if custom_instructions else 'None provided.'}
---
END USER'S CUSTOM INSTRUCTIONS.

BEGIN TRANSCRIPT TEXT:
---
{transcript_text}
---
END TRANSCRIPT TEXT

TASK: Generate the global context summary according to the 'GLOBAL CONTEXT INSTRUCTIONS'. Critically, if specific name spellings (for people, companies, places, etc.) are provided in the 'USER'S CUSTOM INSTRUCTIONS', you MUST use those exact spellings whenever referring to those entities in your generated context. If no custom instructions are provided or they don't mention specific names, proceed as normal based on the transcript."""
        global_context = call_llm_with_retry(context_prompt, selected_provider_context, selected_model_context)
        if not global_context or global_context.startswith("Error:"):
            update_status(15, f"❌ Step 2 Failed: Global context generation error - {global_context}")
            st.stop()
        update_status(25, "✔️ Step 2: Global context generated.")
        with st.expander("View Generated Global Context"):
            st.write(global_context)

        update_status(30, "Step 3/7: Segmenting transcript text...")
        chunk_s = 8000
        chunk_o = 400
        
        # <<< MODIFIED CHUNKING CALL >>>
        text_chunks = enhanced_chunk_text(
            transcript_text,
            provider=selected_provider_context, # Pass provider for topic chunker
            model=selected_model_context,       # Pass model for topic chunker
            chunk_size=chunk_s,
            chunk_overlap=chunk_o
        )
        # <<< END MODIFIED CHUNKING CALL >>>

        if not text_chunks:
            update_status(30, "❌ Step 3 Failed: Failed to segment transcript text after parsing (all methods).")
            st.stop()
        update_status(35, f"✔️ Step 3: Transcript divided into {len(text_chunks)} segments.")

        update_status(40, f"Step 4/7: Summarizing {len(text_chunks)} segments ({selected_model_context})...")
        segment_summaries = []
        total_chunks = len(text_chunks)
        segment_errors = 0
        for i, chunk in enumerate(text_chunks):
            progress_percentage = 40 + int(50 * ((i + 1) / total_chunks))
            update_status(progress_percentage, f"Step 4/7: Summarizing segment {i+1}/{total_chunks}...")
            segment_prompt = f"""BEGIN STANDARD SUMMARY INSTRUCTIONS:
---
{standard_instructions}
---
END STANDARD SUMMARY INSTRUCTIONS.

BEGIN CONTEXT AND DATA FOR SUMMARY:

Overall Case Context:
{global_context}

User's Custom Instructions (Pay close attention to required name spellings):
---
{custom_instructions if custom_instructions else 'None provided.'}
---
END USER'S CUSTOM INSTRUCTIONS.

Transcript Segment To Summarize ({i+1}/{total_chunks}):
--- START SEGMENT ---
{chunk}
--- END SEGMENT ---

END CONTEXT AND DATA FOR SUMMARY.

TASK: Generate a detailed, objective summary for ONLY the 'Transcript Segment To Summarize' provided above. Adhere strictly to the 'STANDARD SUMMARY INSTRUCTIONS'. Use the 'Overall Case Context' and 'User's Custom Instructions' for background and focus, but DO NOT summarize them. Critically, if specific name spellings (for people, companies, places, etc.) are provided in the 'USER'S CUSTOM INSTRUCTIONS', you MUST use those exact spellings whenever referring to those entities in your segment summary. Base your summary *only* on the information present within the 'START SEGMENT' and 'END SEGMENT' markers. Output ONLY the summary text for this segment."""
            segment_summary = call_llm_with_retry(segment_prompt, selected_provider_context, selected_model_context)
            if segment_summary and not segment_summary.startswith("Error:"):
                segment_summaries.append(segment_summary)
            else:
                st.warning(f"Failed to summarize segment {i+1}. Using error placeholder. LLM Response: {segment_summary}")
                segment_summaries.append(f"*** Error summarizing segment {i+1}. See warnings above. ***")
                segment_errors += 1
        if segment_errors > 0:
            update_status(90, f"⚠️ Step 4: Completed segment summarization with {segment_errors} error(s).")
        else:
            update_status(90, f"✔️ Step 4: All {len(segment_summaries)} segments processed successfully.")

        st.markdown("---"); st.subheader("💾 Saving Outputs:")
        base_name = os.path.splitext(uploaded_file.name)[0]
        safe_base_name = "".join(c for c in base_name if c.isalnum() or c in (' ', '.', '_', '-')).strip().replace(' ', '_')
        safe_base_name = safe_base_name or "transcript_summary"
        role_filename_part = deponent_role.lower().replace(' ', '_').replace('/', '_')
        model1_provider_short = selected_provider_context[:3].lower() if selected_provider_context else "unk"
        model1_name_short = selected_model_context.split('/')[-1].split('-')[2].lower() if selected_model_context and '-' in selected_model_context and len(selected_model_context.split('-')) > 2 else selected_model_context.split('/')[-1][:3].lower()
        model1_info = f"{model1_provider_short}-{model1_name_short}"
        model2_provider_short = selected_provider_final[:3].lower() if selected_provider_final else "unk"
        model2_name_short = selected_model_final.split('/')[-1].split('-')[2].lower() if selected_model_final and '-' in selected_model_final and len(selected_model_final.split('-')) > 2 else selected_model_final.split('/')[-1][:3].lower()
        model2_info = f"{model2_provider_short}-{model2_name_short}"

        update_status(92, "Step 5/7: Generating segment summaries document...")
        output1_content = {"Global Context Summary": global_context}
        valid_segment_summary_count = 0
        for i, summary in enumerate(segment_summaries):
            output1_content[f"Segment {i+1} Summary"] = summary
            if "*** Error summarizing segment" not in summary:
                 valid_segment_summary_count += 1
        docx_output1_bytes = generate_docx_bytes_safe(output1_content)
        
        # If DOCX generation failed, try text fallback
        if docx_output1_bytes is None:
            st.warning("DOCX generation failed for segment summaries, saving as text file")
            output1_filename = output1_filename.replace('.docx', '.txt')
            output1_fullpath = os.path.join(OUTPUT_DIRECTORY, output1_filename)
            saved_path1 = save_as_text_fallback(output1_content, output1_fullpath)
        else:
            output1_filename = f"{safe_base_name}_Segments_{model1_info}.docx"
        output1_fullpath = os.path.join(OUTPUT_DIRECTORY, output1_filename)
        saved_path1 = save_bytes_to_file(docx_output1_bytes, output1_fullpath)
        if saved_path1 is None:
            update_status(92, f"❌ Step 5 Failed: Could not save the Segment Summaries DOCX file (Initial target: {output1_filename}).")
            save_success1_flag = False
        else:
            update_status(94, f"✔️ Step 5: Segment summaries document generated (saved as {os.path.basename(saved_path1)}).")
            save_success1_flag = True

        if valid_segment_summary_count == 0:
            update_status(95, "❌ Step 6 Skipped: No valid segment summaries were generated to synthesize.")
            if not save_success1_flag:
                st.error("Processing stopped: No valid segments generated and segment file also failed to save.")
            else:
                st.warning(f"Processing stopped: No valid segments generated, but the segment file (Output 1) was saved as {os.path.basename(saved_path1)}.")
            st.stop()

        update_status(95, f"Step 6/7: Generating final synthesis ({selected_model_final})...")
        combined_summaries_text = "\n\n--- End of Segment / Start of Next ---\n\n".join(
            [f"Segment {i+1} Summary:\n{s}" for i, s in enumerate(segment_summaries) if "*** Error summarizing segment" not in s]
        )
        final_prompt = f"""BEGIN ROLE-SPECIFIC SYNTHESIS INSTRUCTIONS & OUTLINE ({deponent_role}):
---
{role_outline_instructions}
---
END ROLE-SPECIFIC SYNTHESIS INSTRUCTIONS & OUTLINE.

BEGIN DATA TO SYNTHESIZE:

User's Custom Instructions (Pay close attention to required name spellings):
---
{custom_instructions if custom_instructions else 'None provided.'}
---
END USER'S CUSTOM INSTRUCTIONS.

Collection of Segment Summaries:
--- START SUMMARIES ---
{combined_summaries_text}
--- END SUMMARIES ---

END DATA TO SYNTHESIZE.

TASK: You are an experienced civil litigation attorney. Your task is to synthesize the provided 'Collection of Segment Summaries' into a single, cohesive, objective narrative deposition summary. Strictly follow the structure, headings, and requirements outlined in the 'ROLE-SPECIFIC SYNTHESIS INSTRUCTIONS & OUTLINE'. Combine related information from different segments, deduplicate redundant points, and organize the content logically under the specified headings. Use the 'User's Custom Instructions' for focus but ensure the final summary remains objective and based *only* on the provided segment summaries. Critically, if specific name spellings (for people, companies, places, etc.) are provided in the 'USER'S CUSTOM INSTRUCTIONS', you MUST use those exact spellings whenever referring to those entities in your final synthesized summary. Do not invent information. If a required section in the outline cannot be filled from the provided summaries, state that explicitly (e.g., "No testimony regarding X was found in the provided summaries."). Begin your output *directly* with the first heading specified in the outline. Do not include introductory phrases like "Here is the final summary:".
"""
        final_summary = call_llm_with_retry(final_prompt, selected_provider_final, selected_model_final)
        if not final_summary or final_summary.startswith("Error:"):
             update_status(95, f"❌ Step 6 Failed: Final summary generation error - {final_summary}")
             if save_success1_flag:
                 st.warning(f"Final summary generation failed, but the segment summaries document (Output 1) was saved successfully as {os.path.basename(saved_path1)}.")
             else:
                 st.error("Both segment summary saving and final summary generation failed.")
             st.stop()
        update_status(97, "✔️ Step 6: Final cohesive summary generated successfully.")

        update_status(98, "Step 7/7: Generating final summary document...")
        doc_title = f"Final Deposition Summary ({deponent_role}) - {safe_base_name}"
        output2_content = {doc_title: final_summary}
        docx_output2_bytes = generate_docx_bytes_safe(output2_content)
        
        # If DOCX generation failed, try text fallback
        if docx_output2_bytes is None:
            st.warning("DOCX generation failed for final summary, saving as text file")
            output2_filename = output2_filename.replace('.docx', '.txt')
            output2_fullpath = os.path.join(OUTPUT_DIRECTORY, output2_filename)
            saved_path2 = save_as_text_fallback(output2_content, output2_fullpath)
        else:
            output2_filename = f"{safe_base_name}_FinalSummary_{role_filename_part}_{model2_info}.docx"
            output2_fullpath = os.path.join(OUTPUT_DIRECTORY, output2_filename)
            saved_path2 = save_bytes_to_file(docx_output2_bytes, output2_fullpath)
        if saved_path2 is None:
            update_status(98, f"❌ Step 7 Failed: Could not save the Final Summary DOCX file (Initial target: {output2_filename}).")
            save_success2_flag = False
        else:
            update_status(99, f"✔️ Step 7: Final summary document generated (saved as {os.path.basename(saved_path2)}).")
            save_success2_flag = True

        progress_bar.progress(1.0)
        
        # Log processing completion
        processing_time = (datetime.now() - st.session_state.processing_state['processing_start_time']).total_seconds()
        logger.info(f"Processing completed in {processing_time:.1f} seconds")
        
        final_message = ""
        if save_success1_flag and save_success2_flag:
            final_message = f"🎉 Processing complete! Summary files saved in '{OUTPUT_DIRECTORY}' (as '{os.path.basename(saved_path1)}' and '{os.path.basename(saved_path2)}')."
            status_placeholder.success(final_message)
            logger.info(f"Both files saved successfully: {saved_path1}, {saved_path2}")
            st.balloons()
        elif save_success1_flag:
            final_message = f"⚠️ Processing finished, but the Final Summary file (Output 2) failed to save. Segment summaries (Output 1) were saved (as {os.path.basename(saved_path1)}) to '{OUTPUT_DIRECTORY}'. Check logs/permissions."
            status_placeholder.warning(final_message)
        elif save_success2_flag:
             final_message = f"⚠️ Processing finished, but the Segment Summaries file (Output 1) failed to save. Final summary (Output 2) was saved (as {os.path.basename(saved_path2)}) to '{OUTPUT_DIRECTORY}'. Check logs/permissions."
             status_placeholder.warning(final_message)
        else:
            final_message = f"❌ Processing finished, but BOTH summary files failed to save. Check logs and permissions for '{OUTPUT_DIRECTORY}'."
            status_placeholder.error(final_message)

    except Exception as e:
        progress_bar.progress(1.0)
        detailed_error = traceback.format_exc()
        st.error(f"An unexpected error terminated the process: {e}")
        st.error("Traceback:")
        st.code(detailed_error, language='text')
        try: update_status(100, f"❌ Processing stopped due to unexpected error: {e}")
        except Exception: print(f"Failed to update final error status: {e}")

# --- Footer ---
st.markdown("---")
st.markdown("⚠️ **Disclaimer:** AI-generated summaries require careful review by a qualified legal professional. Verify all facts, interpretations, and omissions against the original transcript.")