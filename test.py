# app.py
# Final Version: Multi-Provider, Two-Model Selection, Permissive Google Safety, No Temperature Slider

import streamlit as st
import os
import io  # For handling file uploads and downloads in memory
from dotenv import load_dotenv
import string # For fallback text cleaning in DOCX generation
import traceback # For detailed error logging
import requests # For OpenRouter model fetching

# --- Set page config FIRST ---
st.set_page_config(layout="wide", page_title="Deposition Summarizer")

# --- !!! DEFINE OUTPUT DIRECTORY HERE !!! ---
# Use a raw string (r"...") or escaped backslashes for Windows paths
OUTPUT_DIRECTORY = r"D:\OneDrive\OneDrive - Barnes Maloney PLLC\Documents\Depo Summaries"
# OUTPUT_DIRECTORY = "D:\\OneDrive\\OneDrive - Barnes Maloney PLLC\\Documents\\Depo Summaries" # Alternative

# --- Application Metadata for OpenRouter ---
# Optional: Replace with your actual site URL and app name if you deploy publicly
YOUR_SITE_URL = "http://localhost:8501" # Default Streamlit URL
YOUR_APP_NAME = "Deposition Summarizer"
# ---

# --- Load API Keys from .env ---
load_dotenv()
api_keys = {
    "OpenAI": os.getenv("OPENAI_API_KEY"),
    "Anthropic": os.getenv("ANTHROPIC_API_KEY"),
    "Google": os.getenv("GOOGLE_API_KEY"),
    "OpenRouter": os.getenv("OPENROUTER_API_KEY"),
}

# --- Import Libraries ---
# LLM Libraries
from openai import OpenAI, OpenAIError
from anthropic import Anthropic, APIError as AnthropicAPIError
import google.generativeai as genai
from google.api_core import exceptions as GoogleAPICoreExceptions
from google.generativeai.types import HarmCategory, HarmBlockThreshold

# File Parsing Libraries
import docx # python-docx library
import fitz # PyMuPDF library, often imported as fitz

# Text Processing (Chunking)
from langchain_text_splitters import RecursiveCharacterTextSplitter

# --- Initialize LLM Clients (if keys exist) ---
# This dictionary will hold initialized clients for available providers
clients = {}
client_load_status = {} # To show status in sidebar

# Initialize OpenAI Client
if api_keys["OpenAI"]:
    try:
        clients["OpenAI"] = OpenAI(api_key=api_keys["OpenAI"])
        client_load_status["OpenAI"] = "✅ Initialized"
    except OpenAIError as e:
        client_load_status["OpenAI"] = f"❌ OpenAI Error: {e}"
    except Exception as e:
        client_load_status["OpenAI"] = f"❌ Unexpected Error: {e}"
else:
    client_load_status["OpenAI"] = "⚠️ Key Missing"

# Initialize Anthropic (Claude) Client
if api_keys["Anthropic"]:
    try:
        clients["Anthropic"] = Anthropic(api_key=api_keys["Anthropic"])
        client_load_status["Anthropic"] = "✅ Initialized"
    except AnthropicAPIError as e:
        client_load_status["Anthropic"] = f"❌ Anthropic Error: {e}"
    except Exception as e:
        client_load_status["Anthropic"] = f"❌ Unexpected Error: {e}"
else:
    client_load_status["Anthropic"] = "⚠️ Key Missing"

# Initialize Google (Gemini) Client
if api_keys["Google"]:
    try:
        # Configure the genai library with the API key
        genai.configure(api_key=api_keys["Google"])
        # Store the configured genai module itself as the "client"
        clients["Google"] = genai
        client_load_status["Google"] = "✅ Initialized"
    except GoogleAPICoreExceptions.GoogleAPIError as e:
         client_load_status["Google"] = f"❌ Google API Config Error: {e}"
    except Exception as e: # Catch potential auth errors during configure
        client_load_status["Google"] = f"❌ Google Init Error: {e}"
else:
    client_load_status["Google"] = "⚠️ Key Missing"

# Initialize OpenRouter Client (using OpenAI SDK structure)
if api_keys["OpenRouter"]:
    try:
        clients["OpenRouter"] = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_keys["OpenRouter"],
            # Add headers recommended by OpenRouter
            default_headers={
                 "HTTP-Referer": YOUR_SITE_URL,
                 "X-Title": YOUR_APP_NAME,
            },
        )
        client_load_status["OpenRouter"] = "✅ Initialized"
    except OpenAIError as e:
         client_load_status["OpenRouter"] = f"❌ OpenRouter Config Error: {e}"
    except Exception as e:
         client_load_status["OpenRouter"] = f"❌ Unexpected OpenRouter Init Error: {e}"
else:
    client_load_status["OpenRouter"] = "⚠️ Key Missing"


# --- Model Mapping (Define known models for selection) ---
# This provides default options; OpenRouter list is fetched dynamically.
# Updated OpenAI and Google models as requested.
MODEL_MAPPING = {
    "OpenAI": ["o3-mini", "gpt-4o", "o1-mini"], # Specific models requested
    "Anthropic": [
        "claude-3-7-sonnet-20250219", # Latest Sonnet (as of provided info)
        "claude-3-5-sonnet-20241022", # Latest v2 3.5 Sonnet
        "claude-3-5-haiku-20241022",  # Latest 3.5 Haiku
        "claude-3-opus-20240229",   # Opus still relevant
    ],
    "Google": [
        # List based on user-provided info, using specific IDs where given
        "gemini-2.5-pro-exp-03-25", # Experimental (New Default)
        "gemini-2.0-flash",
        "gemini-2.0-flash-lite",
        "gemini-1.5-flash-8b",
        # Adding the standard 1.5 models with 'models/' prefix for potential broader compatibility
        "models/gemini-1.5-pro",
        "models/gemini-1.5-flash",
    ],
    "OpenRouter": ["Fetching..."] # Placeholder, will be populated by fetch_openrouter_models
}
# Shared cache for fetched OpenRouter models to avoid redundant API calls
OPENROUTER_MODELS_CACHE = None

# --- Helper Functions ---

@st.cache_data(ttl=3600) # Cache fetched models for 1 hour
def fetch_openrouter_models(api_key):
    """Fetches the list of available models from the OpenRouter API."""
    global OPENROUTER_MODELS_CACHE
    if OPENROUTER_MODELS_CACHE is not None and (not OPENROUTER_MODELS_CACHE or not OPENROUTER_MODELS_CACHE[0].startswith("Error:")):
        return OPENROUTER_MODELS_CACHE
    if not api_key: return ["Error: OpenRouter API Key missing"]

    st.info("Fetching OpenRouter models list...")
    try:
        headers = {"Authorization": f"Bearer {api_key}"}
        response = requests.get("https://openrouter.ai/api/v1/models", headers=headers, timeout=20)
        response.raise_for_status()
        models_data = response.json().get("data", [])
        model_ids = sorted([model.get("id") for model in models_data if model.get("id")])
        result = model_ids if model_ids else ["Error: No models found in OpenRouter response"]
        st.info(f"OpenRouter models list fetched ({len(result)} models).")
        OPENROUTER_MODELS_CACHE = result
        return result
    except requests.exceptions.Timeout:
        st.error("Failed to fetch OpenRouter models: Request timed out.")
        result = ["Error: Request Timed Out"]
        OPENROUTER_MODELS_CACHE = result
        return result
    except requests.exceptions.RequestException as e:
        st.error(f"Failed to fetch OpenRouter models: Network error - {e}")
        result = [f"Error: Connection Failed"]
        OPENROUTER_MODELS_CACHE = result
        return result
    except Exception as e:
        st.error(f"Error processing OpenRouter models response: {e}")
        result = [f"Error: Processing Failed"]
        OPENROUTER_MODELS_CACHE = result
        return result

def load_instruction(file_path):
    """Loads text content from a specified instruction file."""
    try:
        normalized_path = os.path.normpath(file_path)
        with open(normalized_path, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        st.warning(f"Instruction file not found: {normalized_path}")
        return ""
    except Exception as e:
        st.error(f"Error reading instruction file {normalized_path}: {e}")
        return ""

# --- MODIFIED call_llm Function ---
# Removed the 'temperature' parameter and its usage in API calls
def call_llm(prompt, provider, model_name):
    """
    Handles the API call to the selected LLM provider and model using default temperature.
    Returns the generated text content or an error string starting with "Error:".
    """
    st.write(f"Attempting LLM call: Provider='{provider}', Model='{model_name}'") # Debugging output (no temp)

    # --- Input Validation ---
    if not provider or not model_name:
        st.error("LLM Call Error: Provider or model name missing.")
        return "Error: Provider or model not specified."

    client = clients.get(provider)
    if not client:
        st.error(f"LLM Call Error: Client for '{provider}' not available. Check API key and initialization status in the sidebar.")
        return f"Error: Client for {provider} not initialized."

    if model_name.startswith("Error:") or model_name == "Fetching...":
         st.error(f"LLM Call Error: Invalid model selected for {provider}: {model_name}")
         return f"Error: Invalid model '{model_name}' selected."

    # --- API Call Logic ---
    try:
        # --- OpenAI and OpenRouter ---
        if provider == "OpenAI" or provider == "OpenRouter":
            response = client.chat.completions.create(
                model=model_name,
                messages=[{"role": "user", "content": prompt}],
                # temperature parameter removed - uses API default
            )
            result = response.choices[0].message.content
            st.write(f"LLM call successful ({provider})")
            return result.strip() if result else ""

        # --- Anthropic (Claude) ---
        elif provider == "Anthropic":
            system_prompt = ""
            user_prompt = prompt
            prompt_lines = prompt.split('\n', 1)
            if len(prompt_lines) > 1 and len(prompt_lines[0]) < 300 and \
               any(keyword in prompt_lines[0].lower() for keyword in ["instructions", "role", "task", "goal"]):
                system_prompt = prompt_lines[0]
                user_prompt = prompt_lines[1]

            max_output_tokens = 4096
            if "claude-3-5" in model_name or "claude-3-7" in model_name:
                max_output_tokens = 8192

            response = client.messages.create(
                model=model_name,
                max_tokens=max_output_tokens,
                system=system_prompt if system_prompt else None,
                messages=[{"role": "user", "content": user_prompt}],
                # temperature parameter removed - uses API default
            )

            if response.stop_reason == "max_tokens":
                 st.warning(f"Anthropic response stopped due to max_tokens ({max_output_tokens}). Output might be truncated.")

            result_text = ""
            if isinstance(response.content, list):
                 result_text = "".join([block.text for block in response.content if hasattr(block, 'text')])
            elif hasattr(response.content, 'text'):
                 result_text = response.content.text
            else:
                 st.warning("Unexpected Anthropic response content structure.")
                 result_text = str(response.content)

            st.write(f"LLM call successful ({provider})")
            return result_text.strip() if result_text else ""

        # --- Google (Gemini) ---
        elif provider == "Google":
            safety_settings = {
                HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
            }

            # Temperature configuration removed
            # generation_config=genai.types.GenerationConfig(temperature=temperature)

            model = client.GenerativeModel(
                model_name=model_name,
                safety_settings=safety_settings,
                # generation_config parameter removed - uses API default temperature
            )

            response = model.generate_content(prompt)

            prompt_feedback = getattr(response, 'prompt_feedback', None)
            block_reason = getattr(prompt_feedback, 'block_reason', None) if prompt_feedback else None

            if block_reason:
                 safety_ratings = getattr(prompt_feedback, 'safety_ratings', [])
                 st.error(f"Google API call blocked. Reason: {block_reason}. Safety Ratings: {safety_ratings}")
                 try: st.warning(f"Full Google Response (Debug): {response}")
                 except Exception: st.warning("Could not display full Google response object.")
                 return f"Error: Google API Blocked Prompt ({block_reason})"

            try:
                result = response.text
            except ValueError as ve:
                 candidates = getattr(response, 'candidates', [])
                 if candidates:
                     st.warning(f"Google response generated candidate(s) but accessing '.text' failed ({ve}). Candidates: {candidates}")
                     if candidates[0].content and candidates[0].content.parts:
                         result = "".join([part.text for part in candidates[0].content.parts if hasattr(part, 'text')])
                         if result: return result.strip()
                 st.error(f"Google API call failed: Could not extract text content. Error: {ve}")
                 return f"Error: Google response missing text content ({ve})."
            except Exception as e_text:
                 st.error(f"Unexpected error accessing Google response text: {e_text}")
                 return f"Error: Failed to access Google response text ({e_text})."

            st.write(f"LLM call successful ({provider})")
            return result.strip() if result else ""

        # --- Unknown Provider ---
        else:
            st.error(f"LLM Call Error: Provider '{provider}' logic not implemented.")
            return f"Error: Unknown provider {provider}."

    # --- Exception Handling (remains the same) ---
    except OpenAIError as e:
        st.error(f"API Error ({provider} via OpenAI SDK): {e}")
        error_detail = str(e).lower()
        if "context_length_exceeded" in error_detail: return f"Error: Context length exceeded for {model_name}."
        if "model_not_found" in error_detail: return f"Error: Model not found: {model_name}."
        if "authentication" in error_detail: return f"Error: Authentication failed for {provider}. Check API Key."
        return f"Error: OpenAI/OR Error - {e}" # More specific return
    except AnthropicAPIError as e:
        st.error(f"API Error (Anthropic): {e}")
        error_detail = str(e).lower()
        if "authentication" in error_detail: return "Error: Authentication failed for Anthropic. Check API Key."
        if "quota" in error_detail or "rate_limit" in error_detail: return f"Error: Anthropic Quota/Rate Limit Exceeded - {e}"
        if "invalid_request" in error_detail: return f"Error: Invalid Request for Anthropic model {model_name} - {e}"
        return f"Error: Anthropic Error - {e}"
    except GoogleAPICoreExceptions.PermissionDenied as e:
         st.error(f"API Error (Google): Permission Denied. Check API Key & API enabled status. Details: {e}")
         return f"Error: Google Permission Denied."
    except GoogleAPICoreExceptions.ResourceExhausted as e:
        st.error(f"API Error (Google): Resource Exhausted (Quota Limit?). Details: {e}")
        return f"Error: Google Quota Exceeded."
    except GoogleAPICoreExceptions.InvalidArgument as e:
         st.error(f"API Error (Google): Invalid Argument. Check model name ('{model_name}')/prompt. Details: {e}")
         # Specific check for model name format (might help user)
         if "model" in str(e).lower() and "not found" in str(e).lower():
             return f"Error: Google Invalid Argument - Model '{model_name}' not found or invalid name format."
         return f"Error: Google Invalid Argument."
    except GoogleAPICoreExceptions.GoogleAPIError as e: # Catch other specific Google errors
        st.error(f"API Error (Google): {e}")
        return f"Error: Google API Error - {e}"
    except requests.exceptions.RequestException as e: # Catch network errors
         st.error(f"Network Error for {provider}: {e}")
         return f"Error: Network Connection Failed ({provider})."
    except Exception as e: # Catch-all for anything else
        st.error(f"An unexpected error occurred during the LLM call ({provider} - {model_name}): {e}")
        st.error(traceback.format_exc())
        return f"Unexpected Error ({provider}) - {e}"


def parse_uploaded_file(uploaded_file):
    """Parses the uploaded file (.txt, .pdf, .docx) and extracts text content."""
    if uploaded_file is None:
        return None, "No file uploaded."

    file_name = uploaded_file.name
    file_type = ""
    if '.' in file_name:
        file_type = file_name.split('.')[-1].lower()

    st.info(f"Parsing file: '{file_name}' (Detected type: '{file_type}')")

    try:
        bytes_data = uploaded_file.getvalue()
        if not bytes_data:
             st.warning(f"Uploaded file '{file_name}' appears to be empty.")
             return "", None

        if file_type == 'pdf':
            text = ""
            is_ocr_needed = False
            try:
                with fitz.open(stream=bytes_data, filetype="pdf") as doc:
                    if not doc.page_count:
                         st.warning(f"PDF file '{file_name}' has zero pages.")
                         return "", None
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        page_text_sorted = page.get_text("text", sort=True).strip()
                        page_text_simple = page.get_text("simple").strip()
                        page_text = page_text_sorted if page_text_sorted else page_text_simple
                        if not page_text and page.get_images(full=True):
                             is_ocr_needed = True
                        if page_text: text += page_text + "\n"
                        if page_num < len(doc) - 1: text += "\n--- Page Break ---\n"
            except fitz.fitz.FileDataError as fe:
                 return None, f"Error parsing PDF '{file_name}': File may be corrupt, password protected, or not a valid PDF format. Details: {fe}"
            except Exception as fitz_e:
                 return None, f"Unexpected error during PDF processing: {fitz_e}\n{traceback.format_exc()}"

            meaningful_text_check = text.replace("--- Page Break ---", "").strip()
            if not meaningful_text_check:
                 if is_ocr_needed:
                      return None, f"Failed to extract text from PDF: '{file_name}'. The document appears to be image-based and requires OCR (Optical Character Recognition), which this tool does not perform."
                 else:
                      st.warning(f"Parsed PDF '{file_name}', but no text content was extracted. The file might be empty or have an unusual structure.")
                      return "", None
            st.success(f"Successfully parsed PDF: '{file_name}'")
            return text, None

        elif file_type == 'docx':
            try:
                doc_io = io.BytesIO(bytes_data)
                doc = docx.Document(doc_io)
                paragraphs_text = [para.text for para in doc.paragraphs if para.text and para.text.strip()]
                text = "\n".join(paragraphs_text)
                if not text.strip():
                     st.warning(f"DOCX file '{file_name}' parsed, but contained no text in its paragraphs.")
                     return "", None
                st.success(f"Successfully parsed DOCX: '{file_name}'")
                return text, None
            except Exception as docx_e:
                 return None, f"Error opening or parsing DOCX file '{file_name}'. It might be corrupted or not a valid DOCX format. Error: {docx_e}\n{traceback.format_exc()}"

        elif file_type == 'txt':
            try:
                text = bytes_data.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    text = bytes_data.decode("cp1252")
                    st.warning("Decoded TXT file using CP1252 encoding (UTF-8 failed). Check for garbled characters.")
                except Exception as e_decode:
                   st.error(f"Failed to decode TXT file '{file_name}' with UTF-8 or CP1252: {e_decode}")
                   return None, f"TXT decode error: Could not decode file content ({e_decode})."
            if not text.strip():
                 st.warning(f"TXT file '{file_name}' parsed, but was empty or contained only whitespace.")
                 return "", None
            st.success(f"Successfully parsed TXT: '{file_name}'")
            return text, None
        else:
            error_msg = f"Unsupported file type: '.{file_type}'. Please upload a file with a .txt, .pdf, or .docx extension."
            st.error(error_msg)
            return None, error_msg
    except Exception as e:
        error_msg = f"An unexpected error occurred during file parsing for '{file_name}': {e}\n{traceback.format_exc()}"
        st.error(error_msg)
        return None, error_msg

def chunk_text(text, chunk_size=8000, chunk_overlap=400):
    """Splits the text into overlapping chunks for LLM processing."""
    st.info(f"Chunking text using RecursiveCharacterTextSplitter (Size: {chunk_size}, Overlap: {chunk_overlap})...")
    if not text or text.isspace():
        st.warning("Input text for chunking is empty or only whitespace.")
        return []
    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size, chunk_overlap=chunk_overlap, length_function=len,
            is_separator_regex=False, separators=["\n\n", "\n", ". ", "! ", "? ", "; ", ": ", ",", " ", ""],
            keep_separator=True
        )
        chunks = text_splitter.split_text(text)
        chunks = [chunk for chunk in chunks if chunk and not chunk.isspace()]
        if not chunks:
            st.warning("Text chunking resulted in zero non-whitespace chunks.")
            if text and not text.isspace():
                st.info("Returning original text as a single chunk.")
                return [text]
            else: return []
        st.write(f"Text divided into {len(chunks)} chunks.")
        return chunks
    except Exception as e:
        st.error(f"An error occurred during text chunking: {e}")
        st.error(traceback.format_exc())
        return []

def generate_docx_bytes(content_dict):
    """Generates a DOCX file's content in memory as bytes from a dictionary."""
    try:
        document = docx.Document()
        for title, text in content_dict.items():
            title_str = str(title) if title is not None else "Untitled Section"
            text_str = str(text) if text is not None else ""
            clean_title = ''.join(filter(lambda x: x in string.printable, title_str)).strip()
            if not clean_title: clean_title = "Untitled Section"
            cleaned_text = ""
            for char in text_str:
                 codepoint = ord(char)
                 if codepoint in (0x9, 0xA, 0xD): cleaned_text += char
                 elif (0x20 <= codepoint <= 0xD7FF) or \
                      (0xE000 <= codepoint <= 0xFFFD) or \
                      (0x10000 <= codepoint <= 0x10FFFF): cleaned_text += char
            try:
                document.add_heading(clean_title, level=1)
                document.add_paragraph(cleaned_text)
            except ValueError as ve:
                st.warning(f"Problem adding content for section '{clean_title}'. Applying basic printable filter as fallback. Error: {ve}")
                fallback_text = ''.join(filter(lambda x: x in string.printable, cleaned_text))
                document.add_paragraph(fallback_text + "\n\n[Content possibly modified due to invalid characters detected by DOCX library]")
        bio = io.BytesIO()
        document.save(bio)
        bio.seek(0)
        return bio.getvalue()
    except Exception as e:
        st.error(f"Error generating DOCX file in memory: {e}")
        st.error(traceback.format_exc())
        return None

def save_bytes_to_file(file_bytes, full_path):
    """Saves byte content (like a generated DOCX) to a specified file path."""
    if file_bytes is None:
        st.error(f"Attempted to save 'None' to {os.path.basename(full_path)}. Document generation likely failed.")
        return False
    try:
        dir_name = os.path.dirname(full_path)
        if dir_name: os.makedirs(dir_name, exist_ok=True)
        with open(full_path, 'wb') as f: f.write(file_bytes)
        st.success(f"✔️ File successfully saved to: {full_path}")
        return True
    except PermissionError:
        st.error(f"Permission denied: Cannot write to the directory '{os.path.dirname(full_path)}'. Please check folder permissions.")
        return False
    except OSError as oe:
        st.error(f"Failed to save file to {full_path}: OS Error - {oe}")
        st.error(traceback.format_exc())
        return False
    except Exception as e:
        st.error(f"An unexpected error occurred while saving file to {full_path}: {e}")
        st.error(traceback.format_exc())
        return False

# --- Load Instructions ---
INSTRUCTIONS_DIR = "instructions"
global_context_instructions_path = os.path.join(INSTRUCTIONS_DIR, "global_context_instructions.txt")
standard_summary_instructions_path = os.path.join(INSTRUCTIONS_DIR, "standard_summary_instructions.txt")
global_context_instructions = load_instruction(global_context_instructions_path)
standard_instructions = load_instruction(standard_summary_instructions_path)
role_files = {
    "Plaintiff": os.path.join(INSTRUCTIONS_DIR, "plaintiff_outline.txt"),
    "Defendant": os.path.join(INSTRUCTIONS_DIR, "defendant_outline.txt"),
    "Expert Witness": os.path.join(INSTRUCTIONS_DIR, "expert_witness_outline.txt"),
    "Fact Witness": os.path.join(INSTRUCTIONS_DIR, "fact_witness_outline.txt"),
}
instruction_load_status = {}
instruction_load_status["Global Context"] = bool(global_context_instructions)
instruction_load_status["Standard Summary"] = bool(standard_instructions)
for role, path in role_files.items():
    instruction_load_status[f"{role} Outline"] = os.path.exists(path)

# --- Sidebar Setup ---
st.sidebar.title("⚙️ Config & Status")
st.sidebar.markdown("---")

available_providers = [p for p, status in client_load_status.items() if "✅" in status]

# Define the NEW default model ID
DEFAULT_GEMINI_MODEL = "gemini-2.5-pro-exp-03-25"

if not available_providers:
      st.sidebar.error("No LLM providers initialized successfully. Check API keys in the `.env` file.")
      selected_provider_context = None
      selected_model_context = None
      selected_provider_final = None
      selected_model_final = None
else:
    # --- LLM Selection for Stage 1: Context & Segments ---
    st.sidebar.subheader("1. Context & Segment LLM")
    st.sidebar.caption("Used for initial context generation and summarizing text chunks.")

    # Determine default provider (prioritize Google if available and has the default model)
    default_provider_index = 0
    if "Google" in available_providers and DEFAULT_GEMINI_MODEL in MODEL_MAPPING["Google"]:
        try:
            default_provider_index = available_providers.index("Google")
        except ValueError: pass # Keep default 0 if Google not found in available list

    selected_provider_context = st.sidebar.selectbox(
        "Select Provider (Context/Segments):",
        options=available_providers,
        index=default_provider_index, # Default to Google if possible, otherwise first available
        key="provider_selector_context",
        help="Choose the LLM service for context generation and segment summaries."
    )

    models_for_provider_context = []
    if selected_provider_context:
        if selected_provider_context == "OpenRouter":
            openrouter_key = api_keys.get("OpenRouter")
            models_for_provider_context = fetch_openrouter_models(openrouter_key) if openrouter_key else ["Error: OpenRouter Key Missing"]
        else:
            models_for_provider_context = MODEL_MAPPING.get(selected_provider_context, ["Error: No models defined"])

        if not isinstance(models_for_provider_context, list): models_for_provider_context = ["Error: Invalid model list format"]

        if models_for_provider_context:
            # --- MODIFIED Default Model Logic (Context) ---
            default_index_context = 0 # Basic fallback
            # Prioritize the requested default Gemini model if provider is Google and model exists
            if selected_provider_context == "Google" and DEFAULT_GEMINI_MODEL in models_for_provider_context:
                 default_index_context = models_for_provider_context.index(DEFAULT_GEMINI_MODEL)
            # Set other sensible defaults for other providers
            elif selected_provider_context == "OpenAI" and "o3-mini" in models_for_provider_context:
                 default_index_context = models_for_provider_context.index("o3-mini")
            elif selected_provider_context == "Anthropic" and "claude-3-5-haiku-20241022" in models_for_provider_context:
                 default_index_context = models_for_provider_context.index("claude-3-5-haiku-20241022")
            elif selected_provider_context == "OpenRouter" and not models_for_provider_context[0].startswith("Error"):
                cheap_or_models = ['google/gemini-flash-1.5', 'openai/gpt-3.5-turbo', 'anthropic/claude-3.5-haiku', 'mistralai/mistral-7b-instruct']
                for m in cheap_or_models:
                     if m in models_for_provider_context: default_index_context = models_for_provider_context.index(m); break
            # Ensure index is valid
            if default_index_context >= len(models_for_provider_context): default_index_context = 0

            selected_model_context = st.sidebar.selectbox(
                 "Select Model (Context/Segments):",
                 options=models_for_provider_context,
                 index=default_index_context,
                 key=f"model_selector_context_{selected_provider_context}",
                 help=f"Choose the specific model from {selected_provider_context} for context generation and segment summarization."
            )
        else:
            st.sidebar.warning(f"No models available or loaded for {selected_provider_context}.")
            selected_model_context = None
    else:
        selected_model_context = None

    # --- LLM Selection for Stage 2: Final Synthesis ---
    st.sidebar.markdown("---")
    st.sidebar.subheader("2. Final Synthesis LLM")
    st.sidebar.caption("Used for combining segment summaries into the final structured output.")

    # Reuse default provider logic from above
    selected_provider_final = st.sidebar.selectbox(
        "Select Provider (Final Synthesis):",
        options=available_providers,
        index=default_provider_index, # Default to Google if possible
        key="provider_selector_final",
        help="Choose the LLM service for the final summary synthesis step."
    )

    models_for_provider_final = []
    if selected_provider_final:
        if selected_provider_final == "OpenRouter":
            openrouter_key = api_keys.get("OpenRouter")
            models_for_provider_final = fetch_openrouter_models(openrouter_key) if openrouter_key else ["Error: OpenRouter Key Missing"]
        else:
            models_for_provider_final = MODEL_MAPPING.get(selected_provider_final, ["Error: No models defined"])

        if not isinstance(models_for_provider_final, list): models_for_provider_final = ["Error: Invalid model list format"]

        if models_for_provider_final:
             # --- MODIFIED Default Model Logic (Final) ---
            default_index_final = 0 # Basic fallback
            # Prioritize the requested default Gemini model if provider is Google and model exists
            if selected_provider_final == "Google" and DEFAULT_GEMINI_MODEL in models_for_provider_final:
                 default_index_final = models_for_provider_final.index(DEFAULT_GEMINI_MODEL)
            # Set other sensible defaults for other providers (often more capable models for final stage)
            elif selected_provider_final == "OpenAI" and "gpt-4o" in models_for_provider_final:
                 default_index_final = models_for_provider_final.index("gpt-4o")
            elif selected_provider_final == "Anthropic" and "claude-3-7-sonnet-20250219" in models_for_provider_final:
                 default_index_final = models_for_provider_final.index("claude-3-7-sonnet-20250219")
            elif selected_provider_final == "OpenRouter" and not models_for_provider_final[0].startswith("Error"):
                capable_or_models = ['openai/gpt-4o', 'anthropic/claude-3.7-sonnet', 'google/gemini-pro-1.5', 'mistralai/mixtral-8x22b-instruct']
                for m in capable_or_models:
                     if m in models_for_provider_final: default_index_final = models_for_provider_final.index(m); break
            # Ensure index is valid
            if default_index_final >= len(models_for_provider_final): default_index_final = 0

            selected_model_final = st.sidebar.selectbox(
                 "Select Model (Final Synthesis):",
                 options=models_for_provider_final,
                 index=default_index_final,
                 key=f"model_selector_final_{selected_provider_final}",
                 help=f"Choose the specific model from {selected_provider_final} for the final synthesis step."
            )
        else:
            st.sidebar.warning(f"No models available or loaded for {selected_provider_final}.")
            selected_model_final = None
    else:
        selected_model_final = None

# --- REMOVED LLM Temperature Setting ---
# st.sidebar.markdown("---")
# st.sidebar.subheader("🌡️ LLM Temperature")
# llm_temperature = st.sidebar.slider(...) # Removed

# --- Display Selected Models & Statuses ---
st.sidebar.markdown("---")
st.sidebar.subheader("Selected LLMs:")
if available_providers:
    if selected_provider_context and selected_model_context and not selected_model_context.startswith("Error:"):
        st.sidebar.write(f"Context/Segments: `{selected_provider_context} / {selected_model_context}`")
    elif selected_provider_context: st.sidebar.warning(f"Context/Segments: Invalid model selected for `{selected_provider_context}`.")
    else: st.sidebar.warning("Context/Segment LLM not selected.")

    if selected_provider_final and selected_model_final and not selected_model_final.startswith("Error:"):
        st.sidebar.write(f"Final Synthesis: `{selected_provider_final} / {selected_model_final}`")
    elif selected_provider_final: st.sidebar.warning(f"Final Synthesis: Invalid model selected for `{selected_provider_final}`.")
    else: st.sidebar.warning("Final Synthesis LLM not selected.")
else: # If no providers available initially
     st.sidebar.warning("No LLM Providers available. Cannot select models.")


# Display API Key & Client Initialization Status
st.sidebar.markdown("---")
st.sidebar.subheader("API Key & Client Status:")
for provider, status in client_load_status.items():
    if "✅" in status: st.sidebar.success(f"{provider}: {status}")
    elif "⚠️" in status: st.sidebar.warning(f"{provider}: {status}")
    else: st.sidebar.error(f"{provider}: {status}")

# Display Instruction File Status
st.sidebar.markdown("---")
st.sidebar.subheader("Instruction Files Status:")
if instruction_load_status["Global Context"]: st.sidebar.success("✅ Global context instructions loaded.")
else: st.sidebar.warning("⚠️ Global context instructions missing or empty.")
if instruction_load_status["Standard Summary"]: st.sidebar.success("✅ Standard summary instructions loaded.")
else: st.sidebar.warning("⚠️ Standard summary instructions missing or empty.")
all_roles_loaded = True
for role, path in role_files.items():
    role_key = f"{role} Outline"
    exists = instruction_load_status.get(role_key, False)
    if exists: st.sidebar.success(f"✅ {role} outline file found.")
    else:
        st.sidebar.error(f"❌ {role} outline file MISSING ({os.path.basename(path)})")
        all_roles_loaded = False
if not all_roles_loaded:
    st.sidebar.error("One or more required role outline files are missing from the 'instructions' folder.")


# --- Streamlit UI (Main Area) ---
st.title("📄 Deposition Transcript Summarizer")
st.markdown(f"""
Upload a transcript (.txt, .pdf, .docx), select the deponent's role, configure **two LLMs** in the sidebar (one for context/segments, one for final synthesis - temperature uses model defaults), add optional instructions, and click generate.

**Output Location:** `{OUTPUT_DIRECTORY}` (Ensure this directory exists and is writable).
""")

# --- Main Area Inputs ---
col1, col2 = st.columns([2, 1])
with col1:
    uploaded_file = st.file_uploader("1. Upload Transcript", type=["txt", "pdf", "docx"], key="file_uploader")
    deponent_role = st.selectbox("2. Select Deponent Role:", options=list(role_files.keys()), key="role_selector")
with col2:
    custom_instructions = st.text_area(
        "3. Optional: Add Custom Instructions", height=180,
        placeholder="E.g., Focus on testimony regarding the 'Acme Contract'. List key entities or specific questions addressed.",
        key="custom_instructions"
    )

# --- Processing Trigger Button ---
st.markdown("---")

process_button_disabled = (
    uploaded_file is None or
    not available_providers or
    selected_provider_context is None or selected_model_context is None or
    selected_model_context.startswith("Error:") or selected_model_context == "Fetching..." or
    selected_provider_final is None or selected_model_final is None or
    selected_model_final.startswith("Error:") or selected_model_final == "Fetching..."
)
process_button_tooltip = "Please upload a transcript file and select valid LLM providers/models for both stages in the sidebar to enable generation." if process_button_disabled else "Start processing the transcript and generate summaries."

process_button = st.button(
    "🚀 Generate & Save Summaries", type="primary", disabled=process_button_disabled,
    key="generate_button", help=process_button_tooltip
)

# --- Main Processing Logic ---
if process_button:
    st.info(f"Starting Process...")
    # Pre-checks
    if not available_providers: st.error("Critical Error: No LLM providers configured."); st.stop()
    if not selected_provider_context or not selected_model_context or selected_model_context.startswith("Error:") or selected_model_context == "Fetching...": st.error("Pre-check Failed: Invalid Context/Segment LLM selection."); st.stop()
    if not selected_provider_final or not selected_model_final or selected_model_final.startswith("Error:") or selected_model_final == "Fetching...": st.error("Pre-check Failed: Invalid Final Synthesis LLM selection."); st.stop()
    if selected_provider_context not in clients or not client_load_status.get(selected_provider_context, "").startswith("✅"): st.error(f"Client initialization error for Context/Segments provider '{selected_provider_context}'. Check API key/status."); st.stop()
    if selected_provider_final not in clients or not client_load_status.get(selected_provider_final, "").startswith("✅"): st.error(f"Client initialization error for Final Synthesis provider '{selected_provider_final}'. Check API key/status."); st.stop()
    if not standard_instructions: st.error("Pre-check Failed: Standard summary instructions missing or empty."); st.stop()
    if not global_context_instructions: st.error("Pre-check Failed: Global context instructions missing or empty."); st.stop()
    selected_role_file = role_files.get(deponent_role)
    if not selected_role_file or not os.path.exists(selected_role_file): st.error(f"Pre-check Failed: Outline file for role '{deponent_role}' is missing."); st.stop()
    role_outline_instructions = load_instruction(selected_role_file)
    if not role_outline_instructions: st.error(f"Pre-check Failed: Outline file for '{deponent_role}' could not be loaded or is empty."); st.stop()
    try:
        os.makedirs(OUTPUT_DIRECTORY, exist_ok=True)
        test_file_path = os.path.join(OUTPUT_DIRECTORY, ".streamlit_permission_test")
        with open(test_file_path, "w") as f_test: f_test.write("test")
        os.remove(test_file_path)
        st.info(f"Output directory checked and seems writable: {OUTPUT_DIRECTORY}")
    except PermissionError:
         st.error(f"Pre-check Failed: Permission denied writing to the output directory: {OUTPUT_DIRECTORY}. Please check folder permissions.")
         st.stop()
    except Exception as e:
        st.error(f"Pre-check Failed: Could not create, access, or write to output directory: {OUTPUT_DIRECTORY}. Error: {e}")
        st.stop()

    # Setup Progress Reporting
    st.markdown("---"); st.subheader("📊 Processing Status:")
    status_placeholder = st.empty(); progress_bar = st.progress(0.0)
    def update_status(percentage, text):
        clamped_percentage_float = max(0.0, min(1.0, percentage / 100.0))
        text_str = str(text) if text is not None else ""
        progress_bar.progress(clamped_percentage_float, text=text_str)
        text_lower = text_str.lower()
        if "error" in text_lower or "failed" in text_lower or "❌" in text_str: status_placeholder.error(text_str)
        elif "warning" in text_lower or "issue" in text_lower or "⚠️" in text_str: status_placeholder.warning(text_str)
        elif "✔️" in text_str: status_placeholder.success(text_str)
        else: status_placeholder.info(text_str)

    # Main Processing Steps
    try:
        # Removed temperature from initial status message
        update_status(0, f"Starting process... Ctx/Seg Model: {selected_model_context}, Final Model: {selected_model_final}")

        # Step 1: Parse Uploaded File
        update_status(5, f"Step 1/7: Parsing '{uploaded_file.name}'...")
        transcript_text, error = parse_uploaded_file(uploaded_file)
        if error: update_status(5, f"❌ Step 1 Failed: File Parsing Error - {error}"); st.stop()
        if transcript_text is None or not transcript_text.strip():
             update_status(5, f"❌ Step 1 Failed: Parsing '{uploaded_file.name}' resulted in empty text content. Cannot proceed.")
             st.stop()
        update_status(10, "✔️ Step 1: File parsed successfully.")

        # Step 2: Generate Global Context
        update_status(15, f"Step 2/7: Generating global context ({selected_model_context})...")
        context_prompt = f"{global_context_instructions}\n\nTRANSCRIPT TEXT:\n---\n{transcript_text}\n---\nEND TRANSCRIPT TEXT"
        # Pass call_llm *without* temperature
        global_context = call_llm(context_prompt, selected_provider_context, selected_model_context)
        if not global_context or global_context.startswith("Error:"):
            update_status(15, f"❌ Step 2 Failed: Global context generation error - {global_context}")
            st.stop()
        update_status(25, "✔️ Step 2: Global context generated.")
        with st.expander("View Generated Global Context"): st.write(global_context)

        # Step 3: Chunk Transcript Text
        update_status(30, "Step 3/7: Segmenting transcript text...")
        chunk_s = 8000; chunk_o = 400
        text_chunks = chunk_text(transcript_text, chunk_size=chunk_s, chunk_overlap=chunk_o)
        if not text_chunks:
            update_status(30, "❌ Step 3 Failed: Failed to segment transcript text after parsing.")
            st.stop()
        update_status(35, f"✔️ Step 3: Transcript divided into {len(text_chunks)} segments.")

        # Step 4: Summarize Segments
        update_status(40, f"Step 4/7: Summarizing {len(text_chunks)} segments ({selected_model_context})...")
        segment_summaries = []; total_chunks = len(text_chunks); segment_errors = 0
        for i, chunk in enumerate(text_chunks):
            progress_percentage = 40 + int(50 * ((i + 1) / total_chunks))
            update_status(progress_percentage, f"Step 4/7: Summarizing segment {i+1}/{total_chunks}...")
            segment_prompt = f"""BEGIN STANDARD SUMMARY INSTRUCTIONS:
---
{standard_instructions}
---
END STANDARD SUMMARY INSTRUCTIONS.

BEGIN CONTEXT AND DATA FOR SUMMARY:

Overall Case Context:
{global_context}

User's Custom Instructions (If any):
{custom_instructions if custom_instructions else 'None provided.'}

Transcript Segment To Summarize ({i+1}/{total_chunks}):
--- START SEGMENT ---
{chunk}
--- END SEGMENT ---

END CONTEXT AND DATA FOR SUMMARY.

TASK: Generate a detailed, objective summary for ONLY the 'Transcript Segment To Summarize' provided above. Adhere strictly to the 'STANDARD SUMMARY INSTRUCTIONS'. Use the 'Overall Case Context' and 'Custom Instructions' for background and focus, but DO NOT summarize them. Base your summary *only* on the information present within the 'START SEGMENT' and 'END SEGMENT' markers. Output ONLY the summary text for this segment."""
            # Pass call_llm *without* temperature
            segment_summary = call_llm(segment_prompt, selected_provider_context, selected_model_context)
            if segment_summary and not segment_summary.startswith("Error:"):
                segment_summaries.append(segment_summary)
            else:
                st.warning(f"Failed to summarize segment {i+1}. Using error placeholder. LLM Response: {segment_summary}")
                segment_summaries.append(f"*** Error summarizing segment {i+1}. See warnings above. ***")
                segment_errors += 1
        if segment_errors > 0: update_status(90, f"⚠️ Step 4: Completed segment summarization with {segment_errors} error(s).")
        else: update_status(90, f"✔️ Step 4: All {len(segment_summaries)} segments processed successfully.")

        # --- Output Generation and Saving ---
        st.markdown("---"); st.subheader("💾 Saving Outputs:")
        base_name = os.path.splitext(uploaded_file.name)[0]
        safe_base_name = "".join(c for c in base_name if c.isalnum() or c in (' ', '.', '_', '-')).strip().replace(' ', '_')
        safe_base_name = safe_base_name or "transcript_summary"
        role_filename_part = deponent_role.lower().replace(' ', '_').replace('/', '_')

        # Shorten model names for filenames (attempt to capture main part)
        def get_short_model_name(provider, model_id):
            if not provider or not model_id: return "unk-unk"
            provider_short = provider[:3].lower()
            name_parts = model_id.split('/')[-1].split('-') # Split by path and dash
            if len(name_parts) > 1:
                # Try to get a meaningful part like 'gpt', 'claude', 'gemini', 'o3'
                core_name = name_parts[0] if name_parts[0] in ['gpt', 'claude', 'gemini'] else name_parts[1] if len(name_parts) > 1 and name_parts[1] in ['mini', 'pro', 'flash', 'haiku', 'sonnet', 'opus'] else name_parts[0][:3]
                return f"{provider_short}-{core_name}"
            return f"{provider_short}-{model_id[:3]}" # Basic fallback

        model1_info = get_short_model_name(selected_provider_context, selected_model_context)
        model2_info = get_short_model_name(selected_provider_final, selected_model_final)

        # Step 5: Generate & Save Output 1 (Context + Segment Summaries)
        update_status(92, "Step 5/7: Generating segment summaries document...")
        output1_content = {"Global Context Summary": global_context}
        valid_segment_summary_count = 0
        for i, summary in enumerate(segment_summaries):
            output1_content[f"Segment {i+1} Summary"] = summary
            if "*** Error summarizing segment" not in summary: valid_segment_summary_count += 1
        docx_output1_bytes = generate_docx_bytes(output1_content)
        output1_filename = f"{safe_base_name}_Segments_{model1_info}.docx"
        output1_fullpath = os.path.join(OUTPUT_DIRECTORY, output1_filename)
        save_success1 = save_bytes_to_file(docx_output1_bytes, output1_fullpath)
        if not save_success1: update_status(92, "❌ Step 5 Failed: Could not save the Segment Summaries DOCX file.")
        else: update_status(94, "✔️ Step 5: Segment summaries document saved successfully.")

        # Step 6: Combine & Refine (Final Synthesis)
        if valid_segment_summary_count == 0:
            update_status(95, "❌ Step 6 Skipped: No valid segment summaries were generated to synthesize.")
            if not save_success1: st.error("Processing stopped: No valid segments generated and segment file also failed to save.")
            else: st.warning("Processing stopped: No valid segments generated, but the segment file (Output 1) was saved.")
            st.stop()

        update_status(95, f"Step 6/7: Generating final synthesis ({selected_model_final})...")
        combined_summaries_text = "\n\n--- End of Segment / Start of Next ---\n\n".join(
            [f"Segment {i+1} Summary:\n{s}" for i, s in enumerate(segment_summaries) if "*** Error summarizing segment" not in s]
        )
        final_prompt = f"""BEGIN ROLE-SPECIFIC SYNTHESIS INSTRUCTIONS & OUTLINE ({deponent_role}):
---
{role_outline_instructions}
---
END ROLE-SPECIFIC SYNTHESIS INSTRUCTIONS & OUTLINE.

BEGIN DATA TO SYNTHESIZE:

User's Custom Instructions (If any):
{custom_instructions if custom_instructions else 'None provided.'}

Collection of Segment Summaries:
--- START SUMMARIES ---
{combined_summaries_text}
--- END SUMMARIES ---

END DATA TO SYNTHESIZE.

TASK: You are an experianced civil litigation attorney. Your task is to synthesize the provided 'Collection of Segment Summaries' into a single, cohesive, objective narrative deposition summary. Strictly follow the structure, headings, and requirements outlined in the 'ROLE-SPECIFIC SYNTHESIS INSTRUCTIONS & OUTLINE'. Combine related information from different segments, deduplicate redundant points, and organize the content logically under the specified headings. Use the 'Custom Instructions' for focus but ensure the final summary remains objective and based *only* on the provided segment summaries. Do not invent information. If a required section in the outline cannot be filled from the provided summaries, state that explicitly (e.g., "No testimony regarding X was found in the provided summaries."). Begin your output *directly* with the first heading specified in the outline. Do not include introductory phrases like "Here is the final summary:".
"""
        # Pass call_llm *without* temperature
        final_summary = call_llm(final_prompt, selected_provider_final, selected_model_final)
        if not final_summary or final_summary.startswith("Error:"):
             update_status(95, f"❌ Step 6 Failed: Final summary generation error - {final_summary}")
             if save_success1: st.warning("Final summary generation failed, but the segment summaries document (Output 1) was saved successfully.")
             else: st.error("Both segment summary saving and final summary generation failed.")
             st.stop()
        update_status(97, "✔️ Step 6: Final cohesive summary generated successfully.")

        # Step 7: Generate & Save Output 2 (Final Summary)
        update_status(98, "Step 7/7: Generating final summary document...")
        doc_title = f"Final Deposition Summary ({deponent_role}) - {safe_base_name}"
        output2_content = {doc_title: final_summary}
        docx_output2_bytes = generate_docx_bytes(output2_content)
        output2_filename = f"{safe_base_name}_FinalSummary_{role_filename_part}_{model2_info}.docx"
        output2_fullpath = os.path.join(OUTPUT_DIRECTORY, output2_filename)
        save_success2 = save_bytes_to_file(docx_output2_bytes, output2_fullpath)
        if not save_success2: update_status(98, "❌ Step 7 Failed: Could not save the Final Summary DOCX file.")
        else: update_status(99, "✔️ Step 7: Final summary document saved successfully.")

        # Final Completion
        progress_bar.progress(1.0)
        if save_success1 and save_success2:
            status_placeholder.success(f"🎉 Processing complete! Both summary files saved in '{OUTPUT_DIRECTORY}'.")
            st.balloons()
        elif save_success1: status_placeholder.warning(f"⚠️ Processing finished, but the Final Summary file (Output 2) failed to save. Segment summaries (Output 1) were saved to '{OUTPUT_DIRECTORY}'. Check logs/permissions.")
        elif save_success2: status_placeholder.warning(f"⚠️ Processing finished, but the Segment Summaries file (Output 1) failed to save. Final summary (Output 2) was saved to '{OUTPUT_DIRECTORY}'. Check logs/permissions.")
        else: status_placeholder.error(f"❌ Processing finished, but BOTH summary files failed to save. Check logs and permissions for '{OUTPUT_DIRECTORY}'.")

    # Catch Unexpected Errors during Main Processing
    except Exception as e:
        progress_bar.progress(1.0)
        detailed_error = traceback.format_exc()
        st.error(f"An unexpected error terminated the process: {e}")
        st.error("Traceback:")
        st.code(detailed_error, language='text')
        try: update_status(100, f"❌ Processing stopped due to unexpected error: {e}")
        except Exception: print(f"Failed to update final error status: {e}")

# --- Footer ---
st.markdown("---")
st.markdown("⚠️ **Disclaimer:** AI-generated summaries require careful review by a qualified legal professional. Verify all facts, interpretations, and omissions against the original transcript.")