# app.py
# Final Version: Integrated Topic-Based Chunking, Enhanced Robustness, and UX Improvements

import streamlit as st
import os
import io
import logging
import time
import traceback
import requests
from dotenv import load_dotenv
import string
import re
from datetime import datetime
from typing import List, Dict

# LLM & File Parsing Libraries
from openai import OpenAI, OpenAIError
from anthropic import Anthropic, APIError as AnthropicAPIError
import google.generativeai as genai
from google.api_core import exceptions as GoogleAPICoreExceptions
from google.generativeai.types import HarmCategory, HarmBlockThreshold
import docx
import fitz
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Import the new, powerful chunker AND its configuration class
# <-- CHANGE: Added ChunkerConfig to the import
from enhanced_topic_chunker_fixed import EnhancedTopicBasedChunker, ChunkerConfig

# --- Set page config FIRST ---
st.set_page_config(layout="wide", page_title="Deposition Summarizer")

# --- Define Constants ---
OUTPUT_DIRECTORY = r"D:\OneDrive\OneDrive - Barnes Maloney PLLC\Documents\Depo Summaries"
YOUR_SITE_URL = "http://localhost:8501"
YOUR_APP_NAME = "Deposition Summarizer"
INSTRUCTIONS_DIR = "instructions"

# --- Setup Logging ---
log_dir = os.path.join(OUTPUT_DIRECTORY, "logs")
os.makedirs(log_dir, exist_ok=True)
log_file = os.path.join(log_dir, f"depo_summarizer_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file, encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)
logger.info(f"Deposition Summarizer started. Output directory: {OUTPUT_DIRECTORY}")

# --- Load API Keys ---
load_dotenv()
api_keys = {
    "OpenAI": os.getenv("OPENAI_API_KEY"),
    "Anthropic": os.getenv("ANTHROPIC_API_KEY"),
    "Google": os.getenv("GOOGLE_API_KEY"),
    "OpenRouter": os.getenv("OPENROUTER_API_KEY"),
    "ChatLLM": os.getenv("CHATLLM_API_KEY"),
}

# --- Initialize LLM Clients ---
clients = {}
client_load_status = {}

# OpenAI
if api_keys["OpenAI"]:
    try:
        clients["OpenAI"] = OpenAI(api_key=api_keys["OpenAI"])
        client_load_status["OpenAI"] = "✅ Initialized"
    except Exception as e:
        client_load_status["OpenAI"] = f"❌ Error: {e}"
else:
    client_load_status["OpenAI"] = "⚠️ Key Missing"

# Anthropic
if api_keys["Anthropic"]:
    try:
        clients["Anthropic"] = Anthropic(api_key=api_keys["Anthropic"])
        client_load_status["Anthropic"] = "✅ Initialized"
    except Exception as e:
        client_load_status["Anthropic"] = f"❌ Error: {e}"
else:
    client_load_status["Anthropic"] = "⚠️ Key Missing"

# Google
if api_keys["Google"]:
    try:
        genai.configure(api_key=api_keys["Google"])
        clients["Google"] = genai
        client_load_status["Google"] = "✅ Initialized"
    except Exception as e:
        client_load_status["Google"] = f"❌ Error: {e}"
else:
    client_load_status["Google"] = "⚠️ Key Missing"

# OpenRouter
if api_keys["OpenRouter"]:
    try:
        clients["OpenRouter"] = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_keys["OpenRouter"],
            default_headers={"HTTP-Referer": YOUR_SITE_URL, "X-Title": YOUR_APP_NAME},
        )
        client_load_status["OpenRouter"] = "✅ Initialized"
    except Exception as e:
        client_load_status["OpenRouter"] = f"❌ Error: {e}"
else:
    client_load_status["OpenRouter"] = "⚠️ Key Missing"

# ChatLLM (Abacus RouteLLM)
if api_keys["ChatLLM"]:
    try:
        clients["ChatLLM"] = OpenAI(
            base_url="https://routellm.abacus.ai/v1",
            api_key=api_keys["ChatLLM"],
        )
        client_load_status["ChatLLM"] = "✅ Initialized"
    except Exception as e:
        client_load_status["ChatLLM"] = f"❌ Error: {e}"
else:
    client_load_status["ChatLLM"] = "⚠️ Key Missing"

# --- Model Mapping ---
MODEL_MAPPING = {
    "OpenAI": ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo"],
    "Anthropic": ["claude-3-5-sonnet-20240620", "claude-3-haiku-20240307", "claude-3-opus-20240229"],
    "Google": ["gemini-1.5-pro-latest", "gemini-1.5-flash-latest"],
    "OpenRouter": ["Fetching..."],
    "ChatLLM": [
        "google/gemini-2.5-flash",
        "anthropic/claude-4-sonnet",
        "route-llm"
    ]
}
OPENROUTER_MODELS_CACHE = None

# --- Helper Functions ---

@st.cache_data(ttl=3600)
def fetch_openrouter_models(api_key):
    global OPENROUTER_MODELS_CACHE
    if OPENROUTER_MODELS_CACHE and not OPENROUTER_MODELS_CACHE[0].startswith("Error:"):
        return OPENROUTER_MODELS_CACHE
    if not api_key: return ["Error: OpenRouter API Key missing"]
    logger.info("Fetching OpenRouter models list...")
    try:
        response = requests.get("https://openrouter.ai/api/v1/models", headers={"Authorization": f"Bearer {api_key}"}, timeout=20)
        response.raise_for_status()
        models_data = response.json().get("data", [])
        model_ids = sorted([model.get("id") for model in models_data if model.get("id")])
        if not model_ids: result = ["Error: No models found"]
        else: result = model_ids
        logger.info(f"Fetched {len(result)} OpenRouter models.")
        OPENROUTER_MODELS_CACHE = result
        return result
    except Exception as e:
        logger.error(f"Failed to fetch OpenRouter models: {e}")
        result = [f"Error: {e}"]
        OPENROUTER_MODELS_CACHE = result
        return result

def load_instruction(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logger.warning(f"Instruction file not found or failed to read: {file_path}. Error: {e}")
        return ""

def call_llm(prompt, provider, model_name):
    # This is the core API call function, now wrapped by call_llm_with_retry
    logger.info(f"LLM Call: Provider='{provider}', Model='{model_name}'")
    client = clients.get(provider)
    if not client: return f"Error: Client for {provider} not initialized."
    if model_name.startswith("Error:") or model_name == "Fetching...": return f"Error: Invalid model '{model_name}' selected."

    try:
        if provider in ["OpenAI", "OpenRouter", "ChatLLM"]:
            response = client.chat.completions.create(model=model_name, messages=[{"role": "user", "content": prompt}])
            return response.choices[0].message.content.strip()

        elif provider == "Anthropic":
            # Enhanced system prompt splitting logic
            system_prompt, user_prompt = "", prompt
            match = re.search(r'BEGIN TRANSCRIPT\s*(?:TEXT|SEGMENT)?\s*:?', prompt, flags=re.IGNORECASE)
            if match:
                system_prompt = prompt[:match.start()].strip()
                user_prompt = prompt[match.start():].strip()
            else:
                user_prompt = prompt

            response = client.messages.create(
                model=model_name,
                max_tokens=8192,
                system=system_prompt,
                messages=[{"role": "user", "content": user_prompt}]
            )
            if response.stop_reason == "max_tokens":
                logger.warning(f"Anthropic response truncated due to max_tokens.")
            return "".join([block.text for block in response.content if hasattr(block, 'text')]).strip()

        elif provider == "Google":
            safety_settings = {category: HarmBlockThreshold.BLOCK_NONE for category in HarmCategory}
            model = client.GenerativeModel(model_name=model_name, safety_settings=safety_settings)
            response = model.generate_content(prompt)
            if response.prompt_feedback.block_reason:
                return f"Error: Google API Blocked Prompt ({response.prompt_feedback.block_reason})"
            return response.text.strip()
        else:
            return f"Error: Unknown provider {provider}."

    except (OpenAIError, AnthropicAPIError, GoogleAPICoreExceptions.GoogleAPIError, requests.exceptions.RequestException) as e:
        logger.error(f"API Error ({provider} - {model_name}): {e}")
        return f"Error: API Call Failed - {e}"
    except Exception as e:
        logger.error(f"Unexpected LLM Error ({provider} - {model_name}): {e}\n{traceback.format_exc()}")
        return f"Error: Unexpected - {e}"

def call_llm_with_retry(prompt, provider, model_name, max_retries=3):
    for attempt in range(max_retries):
        result = call_llm(prompt, provider, model_name)
        if not result.startswith("Error:"):
            return result
        
        # Don't retry auth/not found errors
        if any(keyword in result.lower() for keyword in ["authentication", "not found", "invalid"]):
            logger.error(f"Permanent LLM error, not retrying: {result}")
            return result
            
        logger.warning(f"LLM call failed (attempt {attempt + 1}/{max_retries}): {result}. Retrying...")
        time.sleep(2 ** attempt) # Exponential backoff
    
    logger.error(f"LLM call failed after {max_retries} attempts. Last error: {result}")
    return f"Error: Max retries exceeded. Last error: {result}"

def parse_uploaded_file(uploaded_file):
    file_name = uploaded_file.name
    file_type = file_name.split('.')[-1].lower()
    logger.info(f"Parsing file: '{file_name}' ({file_type})")
    try:
        bytes_data = uploaded_file.getvalue()
        if not bytes_data:
            logger.warning("Uploaded file is empty.")
            return "", None
        
        text = ""
        if file_type == 'pdf':
            with fitz.open(stream=bytes_data, filetype="pdf") as doc:
                text = "".join(page.get_text("text", sort=True) for page in doc)
                if not text.strip() and any(page.get_images(full=True) for page in doc):
                    return None, "PDF appears to be image-based and requires OCR, which is not supported."
        elif file_type == 'docx':
            doc = docx.Document(io.BytesIO(bytes_data))
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        elif file_type == 'txt':
            try:
                text = bytes_data.decode("utf-8")
            except UnicodeDecodeError:
                text = bytes_data.decode("cp1252", errors='replace')
                logger.warning("Decoded TXT with cp1252 (UTF-8 failed).")
        else:
            return None, f"Unsupported file type: '{file_type}'"

        if not text.strip():
            logger.warning(f"Parsed file '{file_name}' but no text content was found.")
            return "", None
            
        logger.info(f"Successfully parsed {len(text)} characters from '{file_name}'.")
        return text, None
    except Exception as e:
        logger.error(f"Failed to parse '{file_name}': {e}\n{traceback.format_exc()}")
        return None, f"Error parsing file: {e}"

def chunk_text(text, chunk_size=8000, chunk_overlap=400):
    """Fallback recursive chunker."""
    logger.info("Using fallback RecursiveCharacterTextSplitter for chunking.")
    if not text or not text.strip(): return []
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_text(text)

def enhanced_chunk_text(text, provider, model, chunk_size=8000, chunk_overlap=400):
    """Primary chunking function using the enhanced topic-based method."""
    if not provider or not model or len(text) < 5000:
        logger.info("Text too short or LLM not specified for topic chunking, using fallback.")
        return chunk_text(text, chunk_size, chunk_overlap)
        
    logger.info(f"Attempting Enhanced Topic-Based Chunking with {provider}/{model}...")
    try:
        # <-- CHANGE: Create the config object first
        chunker_config = ChunkerConfig(
            target_chunk_size=chunk_size,
            overlap_size=chunk_overlap
        )
        
        # <-- CHANGE: Pass the config object to the chunker
        chunker = EnhancedTopicBasedChunker(
            llm_provider=provider,
            llm_model=model,
            config=chunker_config, 
            call_llm_func=call_llm_with_retry
        )
        chunks = chunker.chunk_transcript(text)
        if chunks:
            logger.info(f"Enhanced topic chunking successful: {len(chunks)} chunks created.")
            st.info(f"Topic-based chunking created {len(chunks)} logical segments.")
            return chunks
        else:
            logger.warning("Topic chunking produced no chunks. Falling back.")
            return chunk_text(text, chunk_size, chunk_overlap)
    except Exception as e:
        logger.error(f"Topic-based chunking failed: {e}. Falling back.\n{traceback.format_exc()}")
        st.warning("Topic-based chunking failed, using standard method.")
        return chunk_text(text, chunk_size, chunk_overlap)

def generate_docx_bytes_safe(content_dict: Dict[str, str]) -> bytes:
    """Robust DOCX generation that cleans invalid XML characters."""
    try:
        document = docx.Document()
        for title, text in content_dict.items():
            clean_title = ''.join(c for c in str(title) if c in string.printable) or "Untitled"
            
            # Aggressively clean text for XML compatibility
            clean_text = ""
            for char in str(text):
                codepoint = ord(char)
                if (0x20 <= codepoint <= 0xD7FF) or \
                   (0xE000 <= codepoint <= 0xFFFD) or \
                   (0x10000 <= codepoint <= 0x10FFFF) or \
                   codepoint in (0x9, 0xA, 0xD):
                    clean_text += char
            
            document.add_heading(clean_title, level=1)
            document.add_paragraph(clean_text)

        bio = io.BytesIO()
        document.save(bio)
        bio.seek(0)
        return bio.getvalue()
    except Exception as e:
        logger.error(f"Critical error generating DOCX: {e}")
        return None # Signal failure

def save_as_text_fallback(content_dict, full_path):
    """Saves content as a .txt file if DOCX generation fails."""
    txt_path = full_path.replace('.docx', '.txt')
    logger.warning(f"DOCX generation failed. Saving as fallback text file: {txt_path}")
    try:
        with open(txt_path, 'w', encoding='utf-8') as f:
            for title, content in content_dict.items():
                f.write(f"{title}\n{'='*len(title)}\n\n{content}\n\n")
        return txt_path
    except Exception as e:
        logger.error(f"Failed to save text fallback: {e}")
        return None

def save_bytes_to_file(file_bytes, full_path):
    if not file_bytes: return save_as_text_fallback(st.session_state.get('last_content_dict', {}), full_path)
    
    try:
        dir_name = os.path.dirname(full_path)
        os.makedirs(dir_name, exist_ok=True)
        base, ext = os.path.splitext(full_path)
        counter = 1
        final_path = full_path
        while os.path.exists(final_path):
            final_path = f"{base}_({counter}){ext}"
            counter += 1
        with open(final_path, 'wb') as f:
            f.write(file_bytes)
        logger.info(f"File saved successfully to: {final_path}")
        return final_path
    except Exception as e:
        logger.error(f"Failed to save file '{full_path}': {e}\n{traceback.format_exc()}")
        return None


# --- Load Instructions & Set Up Sidebar ---
global_context_instructions = load_instruction(os.path.join(INSTRUCTIONS_DIR, "global_context_instructions.md"))
standard_instructions = load_instruction(os.path.join(INSTRUCTIONS_DIR, "standard_summary_instructions.md"))
role_files = {
    "Plaintiff": os.path.join(INSTRUCTIONS_DIR, "plaintiff_outline.md"),
    "Defendant": os.path.join(INSTRUCTIONS_DIR, "defendant_outline.md"),
    "Expert Witness": os.path.join(INSTRUCTIONS_DIR, "expert_witness_outline.md"),
    "Fact Witness": os.path.join(INSTRUCTIONS_DIR, "fact_witness_outline.md"),
}

st.sidebar.title("⚙️ Config & Status")
available_providers = [p for p, status in client_load_status.items() if "✅" in status]

if not available_providers:
    st.sidebar.error("No LLM providers initialized. Check `.env` file.")
    selected_provider_context = selected_model_context = selected_provider_final = selected_model_final = None
else:
    # --- LLM Selection ---
    st.sidebar.subheader("1. Context & Segment LLM")
    default_provider_idx = available_providers.index("ChatLLM") if "ChatLLM" in available_providers else (available_providers.index("OpenRouter") if "OpenRouter" in available_providers else 0)
    selected_provider_context = st.sidebar.selectbox("Provider (Context)", available_providers, index=default_provider_idx, key="p1")
    models_context = fetch_openrouter_models(api_keys.get(selected_provider_context)) if selected_provider_context == "OpenRouter" else MODEL_MAPPING.get(selected_provider_context, [])
    default_model_context = next((m for m in ["google/gemini-2.5-flash", "google/gemini-1.5-flash-latest", "claude-3-haiku-20240307", "gpt-4o-mini"] if m in models_context), models_context[0] if models_context else None)
    selected_model_context = st.sidebar.selectbox("Model (Context)", models_context, index=models_context.index(default_model_context) if default_model_context else 0, key="m1")

    st.sidebar.subheader("2. Final Synthesis LLM")
    selected_provider_final = st.sidebar.selectbox("Provider (Final)", available_providers, index=default_provider_idx, key="p2")
    models_final = fetch_openrouter_models(api_keys.get(selected_provider_final)) if selected_provider_final == "OpenRouter" else MODEL_MAPPING.get(selected_provider_final, [])
    default_model_final = next((m for m in ["anthropic/claude-4-sonnet", "anthropic/claude-3.5-sonnet", "openai/gpt-4o", "google/gemini-1.5-pro-latest"] if m in models_final), models_final[0] if models_final else None)
    selected_model_final = st.sidebar.selectbox("Model (Final)", models_final, index=models_final.index(default_model_final) if default_model_final else 0, key="m2")

# --- Status Displays ---
with st.sidebar.expander("API & File Status", expanded=False):
    st.write("**Client Status:**")
    for provider, status in client_load_status.items():
        st.write(f"- {provider}: {status}")
    st.write("**Instruction Files:**")
    st.write(f"- Global Context: {'✅' if global_context_instructions else '❌'}")
    st.write(f"- Standard Summary: {'✅' if standard_instructions else '❌'}")
    for role, path in role_files.items():
        st.write(f"- {role} Outline: {'✅' if os.path.exists(path) else '❌'}")

# --- Main UI ---
st.title("📄 Deposition Transcript Summarizer")
st.markdown(f"Upload a transcript, configure models in the sidebar, and generate a structured summary. Outputs are saved to `{OUTPUT_DIRECTORY}`.")

col1, col2 = st.columns([2, 1])
with col1:
    uploaded_file = st.file_uploader("1. Upload Transcript", type=["txt", "pdf", "docx"])
    deponent_role = st.selectbox("2. Select Deponent Role", list(role_files.keys()))
with col2:
    custom_instructions = st.text_area("3. Optional: Add Custom Instructions (Focus on Names)",
        placeholder="E.g., Johnathan P. Smyth (not Jon Smith)\nACME Corporation Ltd.", height=180)


# --- Process Button & Logic ---
process_button_disabled = not all([uploaded_file, selected_model_context, selected_model_final, not selected_model_context.startswith("Error:"), not selected_model_final.startswith("Error:")])
if st.button("🚀 Generate & Save Summaries", type="primary", disabled=process_button_disabled):
    
    # --- Pre-run Checks ---
    if not all([global_context_instructions, standard_instructions, os.path.exists(role_files[deponent_role])]):
        st.error("One or more essential instruction files are missing. Check status in the sidebar."); st.stop()
    try:
        os.makedirs(OUTPUT_DIRECTORY, exist_ok=True)
        test_file = os.path.join(OUTPUT_DIRECTORY, ".tmp")
        with open(test_file, 'w') as f: f.write('test')
        os.remove(test_file)
    except Exception as e:
        st.error(f"Output directory is not writable: {OUTPUT_DIRECTORY}. Error: {e}"); st.stop()
    
    role_outline_instructions = load_instruction(role_files[deponent_role])

    # --- Main Workflow ---
    status_placeholder = st.empty()
    progress_bar = st.progress(0, "Starting...")

    def update_status(pct, msg):
        progress_bar.progress(pct / 100, msg)
        status_placeholder.info(msg)
        logger.info(msg)

    try:
        # Step 1: Parse File
        update_status(5, f"Step 1/7: Parsing '{uploaded_file.name}'...")
        transcript_text, error = parse_uploaded_file(uploaded_file)
        if error: st.error(error); st.stop()
        if not transcript_text: st.error("Parsing resulted in empty text content."); st.stop()
        update_status(10, "✔️ Step 1: File parsed successfully.")

        # Step 2: Generate Global Context
        update_status(15, f"Step 2/7: Generating global context ({selected_model_context})...")
        context_prompt = f"{global_context_instructions}\n\nCustom Instructions:\n{custom_instructions}\n\nBEGIN TRANSCRIPT TEXT:\n{transcript_text}"
        global_context = call_llm_with_retry(context_prompt, selected_provider_context, selected_model_context)
        if global_context.startswith("Error:"): st.error(f"Global context generation failed: {global_context}"); st.stop()
        update_status(25, "✔️ Step 2: Global context generated.")
        with st.expander("View Generated Global Context"): st.write(global_context)

        # Step 3: Chunk Text
        update_status(30, "Step 3/7: Segmenting transcript into topics...")
        text_chunks = enhanced_chunk_text(transcript_text, selected_provider_context, selected_model_context)
        if not text_chunks: st.error("Failed to segment transcript text."); st.stop()
        update_status(35, f"✔️ Step 3: Transcript divided into {len(text_chunks)} segments.")

        # Step 4: Summarize Segments
        update_status(40, f"Step 4/7: Summarizing {len(text_chunks)} segments...")
        segment_summaries, errors = [], 0
        for i, chunk in enumerate(text_chunks):
            progress = 40 + int(50 * (i + 1) / len(text_chunks))
            update_status(progress, f"Step 4/7: Summarizing segment {i + 1}/{len(text_chunks)}...")
            segment_prompt = f"{standard_instructions}\n\nGlobal Context:\n{global_context}\n\nCustom Instructions:\n{custom_instructions}\n\nBEGIN TRANSCRIPT SEGMENT:\n{chunk}"
            summary = call_llm_with_retry(segment_prompt, selected_provider_context, selected_model_context)
            if summary.startswith("Error:"):
                errors += 1
                segment_summaries.append(f"*** Error summarizing segment {i+1}: {summary} ***")
            else:
                segment_summaries.append(summary)
        
        if errors: update_status(90, f"⚠️ Step 4: Segment summarization completed with {errors} error(s).")
        else: update_status(90, "✔️ Step 4: All segments summarized successfully.")

        # --- Generate & Save Outputs ---
        base_name = os.path.splitext(uploaded_file.name)[0]
        safe_base_name = "".join(c for c in base_name if c.isalnum() or c in (' ', '_', '-')).strip().replace(' ', '_')

        # Output 1: Segment Summaries
        update_status(92, "Step 5/7: Generating segment summaries document...")
        output1_content = {"Global Context": global_context, **{f"Segment {i+1} Summary": s for i, s in enumerate(segment_summaries)}}
        st.session_state['last_content_dict'] = output1_content
        docx_bytes1 = generate_docx_bytes_safe(output1_content)
        path1 = save_bytes_to_file(docx_bytes1, os.path.join(OUTPUT_DIRECTORY, f"{safe_base_name}_Segments.docx"))
        if not path1: update_status(92, "❌ Step 5 Failed: Could not save segment summaries file.")
        else: update_status(94, f"✔️ Step 5: Segment summaries saved.")

        # Step 6: Final Synthesis with Batching Safety
        if not any(s for s in segment_summaries if not s.startswith("*** Error")):
            st.warning("No valid segment summaries were generated. Skipping final synthesis."); st.stop()
        update_status(95, f"Step 6/7: Generating final synthesis ({selected_model_final})...")
        
        # Filter out error summaries
        valid_summaries = [s for s in segment_summaries if not s.startswith("*** Error")]
        combined_summaries = "\n\n---\n\n".join(valid_summaries)
        
        # New: Batch if too long to prevent token limit issues
        max_prompt_chars = 150000  # Conservative for Claude
        if len(combined_summaries) > max_prompt_chars:
            st.info(f"   Large deposition detected ({len(combined_summaries):,} chars), using batched synthesis...")
            batch_size = 5  # Process 5 summaries at a time
            batched_summaries = []
            
            for i in range(0, len(valid_summaries), batch_size):
                batch = valid_summaries[i:i+batch_size]
                if batch:
                    batch_text = "\n\n".join(batch)
                    batch_prompt = f"Summarize these segment summaries concisely, preserving key facts and testimony:\n\n{batch_text}"
                    
                    update_status(95 + int(2 * i / len(valid_summaries)), f"   Processing batch {i//batch_size + 1}/{(len(valid_summaries)-1)//batch_size + 1}...")
                    batch_summary = call_llm_with_retry(batch_prompt, selected_provider_final, selected_model_final)
                    
                    if not batch_summary.startswith("Error:"):
                        batched_summaries.append(batch_summary)
                    else:
                        st.warning(f"Batch {i//batch_size + 1} failed: {batch_summary}")
            
            if batched_summaries:
                combined_summaries = "\n\n---\n\n".join(batched_summaries)
                st.info(f"   ✓ Batched {len(valid_summaries)} summaries into {len(batched_summaries)} intermediate summaries")
            else:
                st.error("All batches failed, cannot proceed with synthesis"); st.stop()
        
        # Final synthesis
        final_prompt = f"{role_outline_instructions}\n\nCustom Instructions:\n{custom_instructions}\n\nBEGIN SUMMARIES TO SYNTHESIZE:\n{combined_summaries}"
        final_summary = call_llm_with_retry(final_prompt, selected_provider_final, selected_model_final)
        if final_summary.startswith("Error:"): st.error(f"Final synthesis failed: {final_summary}"); st.stop()
        update_status(97, "✔️ Step 6: Final summary generated.")

        # Output 2: Final Summary
        update_status(98, "Step 7/7: Generating final summary document...")
        output2_content = {f"Final Summary of {safe_base_name} ({deponent_role})": final_summary}
        st.session_state['last_content_dict'] = output2_content
        docx_bytes2 = generate_docx_bytes_safe(output2_content)
        path2 = save_bytes_to_file(docx_bytes2, os.path.join(OUTPUT_DIRECTORY, f"{safe_base_name}_FinalSummary.docx"))
        if not path2: update_status(98, "❌ Step 7 Failed: Could not save final summary file.")
        else: update_status(99, f"✔️ Step 7: Final summary saved.")
        
        # Completion
        progress_bar.progress(100)
        if path1 and path2:
            status_placeholder.success(f"🎉 Processing complete! Files saved to '{OUTPUT_DIRECTORY}'.")
            st.balloons()
        else:
            status_placeholder.warning("Processing finished, but one or more files failed to save. Check logs.")
            
    except Exception as e:
        logger.error(f"A critical error terminated the process: {e}\n{traceback.format_exc()}")
        st.error(f"An unexpected error terminated the process: {e}")
        status_placeholder.error(f"❌ Processing stopped unexpectedly. Check logs for details.")

# --- Footer ---
st.markdown("---")
st.markdown("⚠️ **Disclaimer:** AI-generated summaries require careful review by a qualified legal professional. Verify all facts, interpretations, and omissions against the original transcript.")